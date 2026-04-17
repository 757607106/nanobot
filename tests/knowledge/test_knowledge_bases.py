from __future__ import annotations

import asyncio
import io
import time
from pathlib import Path

import pytest

from nanobot.config.schema import Config
from nanobot.platform.instances import PlatformInstance
from nanobot.platform.knowledge import (
    KnowledgeBaseService,
    KnowledgeBaseValidationError,
    create_knowledge_store,
)
from nanobot.platform.knowledge.preview_artifacts import KnowledgePreviewArtifacts
from nanobot.platform.knowledge.rag_engine import IndexResult
from tests.knowledge_test_utils import FakeRAGEngine


def _make_instance(tmp_path: Path) -> PlatformInstance:
    unique_id = f"instance-{tmp_path.parent.name}-{tmp_path.name}"
    return PlatformInstance(
        id=unique_id,
        label="Test Instance",
        config_path=tmp_path / "instance" / "config.json",
    )


def _make_store(instance: PlatformInstance, config: Config | None = None):
    return create_knowledge_store(config or Config(), instance)


def _wait_for_job(service: KnowledgeBaseService, kb_id: str, job_id: str) -> dict[str, object]:
    deadline = time.time() + 5.0
    while time.time() < deadline:
        job = next((item for item in service.list_jobs(kb_id) if item["jobId"] == job_id), None)
        if job and job["status"] in {"succeeded", "failed"}:
            return job
        time.sleep(0.05)
    raise AssertionError(f"Timed out waiting for job {job_id}")


def _build_docx_bytes(*paragraphs: str) -> bytes:
    from docx import Document

    buffer = io.BytesIO()
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(buffer)
    return buffer.getvalue()


def _build_xlsx_bytes(*rows: tuple[object, ...]) -> bytes:
    from openpyxl import Workbook

    buffer = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(list(row))
    workbook.save(buffer)
    return buffer.getvalue()


class _RuntimeAwareFakeRAGEngine(FakeRAGEngine):
    def __init__(self) -> None:
        super().__init__()
        self.reset_calls: list[str] = []

    async def reset_kb(self, kb_id: str) -> bool:
        self.reset_calls.append(kb_id)
        return True


class _SlowQueryFakeRAGEngine(FakeRAGEngine):
    async def query_structured(self, kb_id, query_text, **kwargs):
        await asyncio.sleep(0.2)
        return await super().query_structured(kb_id, query_text, **kwargs)


class _MultimodalRoutingFakeRAGEngine(FakeRAGEngine):
    def __init__(self) -> None:
        super().__init__()
        self.raw_file_calls: list[dict[str, str | None]] = []
        self.text_calls: list[dict[str, str | None]] = []

    async def insert_text(
        self,
        kb_id: str,
        text: str,
        *,
        doc_id: str | None = None,
        file_path: str | None = None,
    ) -> IndexResult:
        self.text_calls.append(
            {
                "kb_id": kb_id,
                "doc_id": doc_id,
                "file_path": file_path,
            }
        )
        return await super().insert_text(kb_id, text, doc_id=doc_id, file_path=file_path)

    async def insert_document_file(
        self,
        kb_id: str,
        file_path: str,
        *,
        doc_id: str | None = None,
        file_name: str | None = None,
    ) -> IndexResult:
        self.raw_file_calls.append(
            {
                "kb_id": kb_id,
                "doc_id": doc_id,
                "file_path": file_path,
                "file_name": file_name,
            }
        )
        stored_doc_id = str(doc_id or file_name or file_path)
        self._docs.setdefault(kb_id, {})[stored_doc_id] = {
            "content": f"raw:{file_name or Path(file_path).name}",
            "file_path": str(file_name or Path(file_path).name),
            "doc_id": stored_doc_id,
            "chunks": [f"raw:{file_name or Path(file_path).name}"],
        }
        return IndexResult(
            success=True,
            doc_id=stored_doc_id,
            chunks_count=3,
            track_id=f"track-{stored_doc_id}",
        )


def test_knowledge_base_service_file_and_source_flow(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )

    created = service.create_knowledge_base(
        {
            "name": "Ops Handbook",
            "description": "Runbooks and operating notes",
            "kbType": "lightrag",
        }
    )
    kb_id = created["kbId"]

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "handover.md",
                "mime_type": "text/markdown",
                "content": b"# Handover\n\nEscalation path: page the on-call engineer before restarting shared services.\n",
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]

    faq_source = service.add_source_file(
        kb_id,
        {
            "sourceType": "faq_table",
            "title": "Ops FAQ",
            "items": [
                {
                    "question": "How do we restart nanobot?",
                    "answer": "Use supervisorctl restart nanobot after checking the current process state.",
                }
            ],
        },
    )

    parse_job = service.parse_files(kb_id, {"file_ids": [file_id, faq_source["fileId"]]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, parse_job)["status"] == "succeeded"

    index_job = service.index_files(kb_id, {"file_ids": [file_id, faq_source["fileId"]]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, index_job)["status"] == "succeeded"

    listed_files = service.list_files(kb_id)
    assert len(listed_files["items"]) == 2
    assert all(item["status"] == "indexed" for item in listed_files["items"])

    retrieved = service.retrieve(
        kb_ids=[kb_id],
        query="restart nanobot service",
        limit=4,
    )
    assert len(retrieved["hits"]) >= 1
    assert any("supervisorctl restart nanobot" in hit["content"] for hit in retrieved["hits"])


def test_knowledge_base_service_delete_files(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )

    created = service.create_knowledge_base({"name": "Support KB"})
    kb_id = created["kbId"]

    first = service.upload_files(
        kb_id,
        [
            {
                "file_name": "runbook.md",
                "mime_type": "text/markdown",
                "content": b"# Runbook\n\nRestart the worker.\n",
            }
        ],
    )
    second = service.upload_files(
        kb_id,
        [
            {
                "file_name": "faq.md",
                "mime_type": "text/markdown",
                "content": b"# FAQ\n\nReset the token cache.\n",
            }
        ],
    )

    file_ids = [first["items"][0]["fileId"], second["items"][0]["fileId"]]
    deleted = service.delete_files(kb_id, file_ids)
    assert deleted == {"deleted_count": 2, "file_ids": file_ids}
    assert service.list_files(kb_id)["items"] == []
    assert service.list_jobs(kb_id) == []


def test_query_kb_for_agent_uses_fast_context_only_mode(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )

    created = service.create_knowledge_base({"name": "Agent KB"})
    kb_id = created["kbId"]

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "faq.md",
                "mime_type": "text/markdown",
                "content": b"Restart nanobot with supervisorctl restart nanobot.\n",
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]
    parse_job = service.parse_files(kb_id, {"file_ids": [file_id]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, parse_job)["status"] == "succeeded"
    index_job = service.index_files(kb_id, {"file_ids": [file_id]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, index_job)["status"] == "succeeded"

    result = service.query_kb_for_agent(kb_id, "How do we restart nanobot?", limit=4)

    assert result["metadata"]["mode"] == "mix"
    assert result["query_params"]["only_need_context"] is True
    assert result["query_params"]["top_k"] == 4


def test_query_database_honors_internal_best_effort_timeout(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=_SlowQueryFakeRAGEngine(),
    )

    created = service.create_knowledge_base({"name": "Slow KB"})
    kb_id = created["kbId"]

    with pytest.raises(TimeoutError, match="Knowledge async task timed out"):
        service.query_database(
            kb_id,
            {
                "query": "hello",
                "__best_effort_timeout_seconds__": 0.01,
            },
        )


def test_query_database_routes_multimodal_requests_through_rag_query(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    rag_engine = FakeRAGEngine()
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=rag_engine,
    )

    kb_id = service.create_knowledge_base({"name": "Multimodal Query KB"})["kbId"]
    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "diagram.png",
                "mime_type": "image/png",
                "content": b"\x89PNG\r\n\x1a\nquery-image",
            }
        ],
    )
    image_file_id = uploaded["items"][0]["fileId"]

    result = service.query_database(
        kb_id,
        {
            "query": "Describe the uploaded image",
            "mode": "mix",
            "only_need_context": False,
            "multimodal_content": [
                {
                    "type": "image",
                    "file_id": image_file_id,
                }
            ],
        },
    )

    assert result["metadata"]["multimodalQuery"] is True
    assert result["metadata"]["multimodalContentTypes"] == ["image"]
    assert result["data"]["chunks"] == []
    assert "multimodal[image]" in str(result.get("message") or "")
    assert len(rag_engine.multimodal_queries) == 1
    resolved_content = rag_engine.multimodal_queries[0]["multimodal_content"][0]
    assert resolved_content["file_id"] == image_file_id
    assert str(resolved_content["img_path"]).endswith(".png")


def test_query_database_rejects_non_image_file_for_image_multimodal_query(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )

    kb_id = service.create_knowledge_base({"name": "Multimodal Guard KB"})["kbId"]
    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "notes.md",
                "mime_type": "text/markdown",
                "content": b"# Notes\n",
            }
        ],
    )
    text_file_id = uploaded["items"][0]["fileId"]

    with pytest.raises(KnowledgeBaseValidationError, match="type='image'"):
        service.query_database(
            kb_id,
            {
                "query": "Describe the uploaded image",
                "multimodal_content": [
                    {
                        "type": "image",
                        "file_id": text_file_id,
                    }
                ],
            },
        )


def test_office_preview_variant_generates_html_artifacts(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )
    kb_id = service.create_knowledge_base({"name": "Office Preview KB"})["kbId"]

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "manual.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": _build_docx_bytes("D9 操作手册", "第一步：检查配置。"),
            },
            {
                "file_name": "metrics.xlsx",
                "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "content": _build_xlsx_bytes(
                    ("Name", "Value"),
                    ("CPU", 42),
                    ("Memory", "7GB"),
                ),
            },
        ],
    )
    docx_id = next(item["fileId"] for item in uploaded["items"] if item["filename"] == "manual.docx")
    xlsx_id = next(item["fileId"] for item in uploaded["items"] if item["filename"] == "metrics.xlsx")

    docx_preview = service.get_file_preview(kb_id, docx_id)
    xlsx_preview = service.get_file_preview(kb_id, xlsx_id)
    assert docx_preview["previewKind"] == "html"
    assert docx_preview["contentType"] == "text/html"
    assert xlsx_preview["previewKind"] == "html"
    assert xlsx_preview["contentType"] == "text/html"

    docx_artifact = service.get_file_artifact(kb_id, docx_id, variant="preview")
    xlsx_artifact = service.get_file_artifact(kb_id, xlsx_id, variant="preview")
    assert docx_artifact.path.suffix == ".html"
    assert xlsx_artifact.path.suffix == ".html"
    assert "D9 操作手册" in docx_artifact.path.read_text(encoding="utf-8")
    assert "CPU" in xlsx_artifact.path.read_text(encoding="utf-8")

    docx_record = service.store.get_file(docx_id)
    xlsx_record = service.store.get_file(xlsx_id)
    assert docx_record is not None
    assert xlsx_record is not None
    assert "previewArtifacts" in dict(docx_record.processing_params)
    assert "previewArtifacts" in dict(xlsx_record.processing_params)


def test_delete_file_cleans_generated_office_preview_artifact(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )
    kb_id = service.create_knowledge_base({"name": "Office Preview Cleanup KB"})["kbId"]

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "cleanup.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": _build_docx_bytes("cleanup content"),
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]
    artifact = service.get_file_artifact(kb_id, file_id, variant="preview")
    generated_path = artifact.path
    assert generated_path.exists()

    deleted = service.delete_file(kb_id, file_id)
    assert deleted is True
    assert not generated_path.exists()


def test_presentation_preview_variant_generates_pdf_artifact_and_reuses_cache(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )
    kb_id = service.create_knowledge_base({"name": "Presentation Preview KB"})["kbId"]
    convert_calls = {"count": 0}

    def _fake_convert(_source_path: Path, _preview_dir: Path, target_path: Path) -> tuple[Path | None, str | None]:
        convert_calls["count"] += 1
        target_path.write_bytes(b"%PDF-1.7\n%nanobot-preview\n")
        return target_path, None

    monkeypatch.setattr(
        KnowledgePreviewArtifacts,
        "_convert_presentation_to_pdf",
        staticmethod(_fake_convert),
    )

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "deck.pptx",
                "mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "content": b"pptx-placeholder",
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]

    preview = service.get_file_preview(kb_id, file_id)
    assert preview["previewKind"] == "pdf"
    assert preview["contentType"] == "application/pdf"

    first = service.get_file_artifact(kb_id, file_id, variant="preview")
    second = service.get_file_artifact(kb_id, file_id, variant="preview")
    assert first.path.suffix == ".pdf"
    assert first.filename.endswith(".pdf")
    assert second.path == first.path
    assert second.filename.endswith(".pdf")
    assert convert_calls["count"] == 1

    record = service.store.get_file(file_id)
    assert record is not None
    office_preview = (
        dict(record.processing_params)
        .get("previewArtifacts", {})
        .get("officePreview", {})
    )
    assert office_preview.get("status") == "ready"
    assert office_preview.get("previewKind") == "pdf"


def test_presentation_preview_unavailable_is_cached_without_repeat_conversion(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )
    kb_id = service.create_knowledge_base({"name": "Presentation Preview Fallback KB"})["kbId"]
    convert_calls = {"count": 0}

    def _unavailable_convert(_source_path: Path, _preview_dir: Path, _target_path: Path) -> tuple[Path | None, str | None]:
        convert_calls["count"] += 1
        return None, "soffice unavailable"

    monkeypatch.setattr(
        KnowledgePreviewArtifacts,
        "_convert_presentation_to_pdf",
        staticmethod(_unavailable_convert),
    )

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "fallback.pptx",
                "mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "content": b"pptx-placeholder",
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]

    first_preview = service.get_file_preview(kb_id, file_id)
    assert first_preview["previewKind"] == "unsupported"
    assert first_preview["contentType"].startswith("application/vnd.openxmlformats-officedocument")
    assert convert_calls["count"] == 1

    artifact = service.get_file_artifact(kb_id, file_id, variant="preview")
    assert artifact.path.suffix == ".pptx"

    second_preview = service.get_file_preview(kb_id, file_id)
    assert second_preview["previewKind"] == "unsupported"
    assert convert_calls["count"] == 1

    record = service.store.get_file(file_id)
    assert record is not None
    office_preview = (
        dict(record.processing_params)
        .get("previewArtifacts", {})
        .get("officePreview", {})
    )
    assert office_preview.get("status") == "unavailable"
    assert "soffice unavailable" in str(office_preview.get("reason") or "")


def test_docx_preview_without_mammoth_is_cached_as_unavailable(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )
    kb_id = service.create_knowledge_base({"name": "Docx Preview Dependency Guard KB"})["kbId"]

    def _missing_mammoth(_path: Path) -> str:
        raise RuntimeError("docx preview converter unavailable: mammoth unavailable")

    monkeypatch.setattr(
        KnowledgePreviewArtifacts,
        "_docx_to_html",
        staticmethod(_missing_mammoth),
    )

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "manual.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": _build_docx_bytes("D9 手册"),
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]

    first_preview = service.get_file_preview(kb_id, file_id)
    assert first_preview["previewKind"] == "unsupported"
    assert first_preview["contentType"].startswith("application/vnd.openxmlformats-officedocument")

    second_preview = service.get_file_preview(kb_id, file_id)
    assert second_preview["previewKind"] == "unsupported"

    record = service.store.get_file(file_id)
    assert record is not None
    office_preview = (
        dict(record.processing_params)
        .get("previewArtifacts", {})
        .get("officePreview", {})
    )
    assert office_preview.get("status") == "unavailable"
    assert "mammoth unavailable" in str(office_preview.get("reason") or "")


def test_knowledge_base_service_resolves_kb_model_bindings_for_rag_runtime(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    rag_engine = _RuntimeAwareFakeRAGEngine()
    config = Config.model_validate(
        {
            "providers": {
                "deepseek": {
                    "apiKey": "sk-llm",
                    "apiBase": "https://api.deepseek.com",
                },
                "dashscope": {
                    "apiKey": "sk-embed",
                    "apiBase": "https://dashscope.aliyuncs.com/compatible-mode/v1",
                },
            },
            "modelBindings": {
                "deepseek": {
                    "provider": "deepseek",
                    "label": "DeepSeek Chat",
                    "model": "deepseek-chat",
                    "apiKey": "sk-llm",
                    "apiBase": "https://api.deepseek.com",
                    "capabilityType": "text_chat",
                },
                "text-embedding-v4-2": {
                    "provider": "dashscope",
                    "label": "text-embedding-v4",
                    "model": "text-embedding-v4",
                    "apiKey": "sk-embed",
                    "apiBase": "https://dashscope.aliyuncs.com/compatible-mode/v1",
                    "capabilityType": "embedding",
                },
                "vision-main": {
                    "provider": "openai",
                    "label": "Vision Main",
                    "model": "gpt-4o",
                    "apiKey": "sk-vision",
                    "apiBase": "https://api.openai.com/v1",
                    "capabilityType": "multimodal",
                },
            },
            "rag": {
                "lightragBaseUrl": "http://127.0.0.1:9621",
                "visionBinding": "vision-main",
            },
        }
    )
    service = KnowledgeBaseService(
        _make_store(instance, config),
        instance=instance,
        instance_id=instance.id,
        rag_engine=rag_engine,
        config=config,
    )

    created = service.create_knowledge_base(
        {
            "name": "Binding KB",
            "embedInfo": {
                "bindingName": "text-embedding-v4-2",
                "modelName": "text-embedding-v4",
            },
            "llmInfo": {
                "bindingName": "deepseek",
                "modelName": "deepseek-chat",
            },
            "additionalParams": {
                "enable_multimodal": True,
            },
        }
    )

    # After refactor, RAG engine is an HTTP client — no runtime resolver needed.
    # We just verify the service accepted the model binding config.
    assert created["llmInfo"]["bindingName"] == "deepseek"
    assert created["llmInfo"]["modelName"] == "deepseek-chat"
    assert created["embedInfo"]["bindingName"] == "text-embedding-v4-2"
    assert created["embedInfo"]["modelName"] == "text-embedding-v4"
    runtime = service._resolve_kb_runtime_overrides(str(created["kbId"]))  # noqa: SLF001 - validate KB runtime mapping
    assert runtime["llm_model"] == "deepseek-chat"
    assert runtime["embedding_model"] == "text-embedding-v4"
    assert runtime["embedding_dim"] == 1024
    assert runtime["vision_model"] == "gpt-4o"


def test_knowledge_base_service_blocks_embedding_change_when_indexed(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    rag_engine = _RuntimeAwareFakeRAGEngine()
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=rag_engine,
    )

    created = service.create_knowledge_base({"name": "Immutable Embed KB"})
    kb_id = str(created["kbId"])

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "faq.md",
                "mime_type": "text/markdown",
                "content": b"Restart nanobot with supervisorctl restart nanobot.\n",
            }
        ],
    )
    file_id = uploaded["items"][0]["fileId"]
    parse_job = service.parse_files(kb_id, {"file_ids": [file_id]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, parse_job)["status"] == "succeeded"
    index_job = service.index_files(kb_id, {"file_ids": [file_id]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, index_job)["status"] == "succeeded"

    with pytest.raises(KnowledgeBaseValidationError):
        service.update_knowledge_base(
            kb_id,
            {
                "embedInfo": {
                    "bindingName": "text-embedding-v4-2",
                    "modelName": "text-embedding-v4",
                },
            },
        )


def test_knowledge_base_store_initializes_current_schema(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    store = _make_store(instance)

    with store._connection() as conn:  # noqa: SLF001 - integration-level schema verification
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT tablename
                FROM pg_tables
                WHERE schemaname = 'public'
                AND tablename IN ('knowledge_bases', 'knowledge_files', 'knowledge_jobs')
                """
            )
            tables = {str(row["tablename"]) for row in cur.fetchall()}

            cur.execute(
                """
                SELECT indexname
                FROM pg_indexes
                WHERE schemaname = 'public'
                AND indexname IN ('idx_knowledge_files_kb', 'idx_knowledge_jobs_kb')
                """
            )
            indexes = {str(row["indexname"]) for row in cur.fetchall()}

    assert {"knowledge_bases", "knowledge_files", "knowledge_jobs"} <= tables
    assert {"idx_knowledge_files_kb", "idx_knowledge_jobs_kb"} <= indexes


def test_knowledge_base_service_uses_smaller_default_chunks(tmp_path: Path) -> None:
    instance = _make_instance(tmp_path)
    service = KnowledgeBaseService(
        _make_store(instance),
        instance=instance,
        instance_id=instance.id,
        rag_engine=FakeRAGEngine(),
    )

    kb_payload = service.create_knowledge_base({"name": "Chunk Defaults"})
    kb = service.store.get_kb(str(kb_payload["kbId"]))
    assert kb is not None

    uploaded = service.upload_files(
        str(kb_payload["kbId"]),
        [
            {
                "file_name": "long.md",
                "mime_type": "text/markdown",
                "content": b"placeholder",
            }
        ],
    )
    file_record = service.store.get_file(str(uploaded["items"][0]["fileId"]))
    assert file_record is not None

    chunk_texts = service._build_chunk_texts(kb, file_record, "A" * 900)

    assert len(chunk_texts) > 1
    assert max(len(item) for item in chunk_texts) <= 500


def test_knowledge_base_service_routes_multimodal_docs_to_raw_file_ingest(tmp_path: Path) -> None:
    docx_module = pytest.importorskip("docx")
    openpyxl_module = pytest.importorskip("openpyxl")

    instance = _make_instance(tmp_path)
    rag_engine = _MultimodalRoutingFakeRAGEngine()
    config = Config.model_validate(
        {
            "providers": {
                "openai": {
                    "apiKey": "sk-vision",
                    "apiBase": "https://api.openai.com/v1",
                },
            },
            "modelBindings": {
                "vision-main": {
                    "provider": "openai",
                    "label": "Vision Main",
                    "model": "gpt-4o",
                    "apiKey": "sk-vision",
                    "apiBase": "https://api.openai.com/v1",
                    "capabilityType": "multimodal",
                },
            },
        }
    )
    service = KnowledgeBaseService(
        _make_store(instance, config),
        instance=instance,
        instance_id=instance.id,
        rag_engine=rag_engine,
        config=config,
    )

    created = service.create_knowledge_base(
        {
            "name": "Multimodal KB",
            "additionalParams": {
                "enable_multimodal": True,
                "visionInfo": {
                    "bindingName": "vision-main",
                    "modelName": "gpt-4o",
                },
            },
        }
    )
    kb_id = created["kbId"]

    doc = docx_module.Document()
    doc.add_paragraph("Quarterly report")
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)

    workbook = openpyxl_module.Workbook()
    worksheet = workbook.active
    worksheet.append(["Metric", "Value"])
    worksheet.append(["CPU", 87])
    xlsx_buffer = io.BytesIO()
    workbook.save(xlsx_buffer)

    uploaded = service.upload_files(
        kb_id,
        [
            {
                "file_name": "report.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": docx_buffer.getvalue(),
            },
            {
                "file_name": "metrics.xlsx",
                "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "content": xlsx_buffer.getvalue(),
            },
            {
                "file_name": "deck.pptx",
                "mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "content": b"pptx-placeholder",
            },
            {
                "file_name": "topology.png",
                "mime_type": "image/png",
                "content": b"png-placeholder",
            },
            {
                "file_name": "notes.md",
                "mime_type": "text/markdown",
                "content": b"# Notes\n\nRestart nanobot after deployment.\n",
            },
        ],
    )
    items_by_name = {item["filename"]: item for item in uploaded["items"]}

    parse_job = service.parse_files(kb_id, {"file_ids": [item["fileId"] for item in uploaded["items"]]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, parse_job)["status"] == "succeeded"

    index_job = service.index_files(kb_id, {"file_ids": [item["fileId"] for item in uploaded["items"]]})["job"]["jobId"]
    assert _wait_for_job(service, kb_id, index_job)["status"] == "succeeded"

    raw_file_names = {str(item["file_name"]) for item in rag_engine.raw_file_calls}
    assert raw_file_names == {"report.docx", "metrics.xlsx", "deck.pptx", "topology.png"}
    assert len(rag_engine.text_calls) == 1
    assert rag_engine.text_calls[0]["doc_id"] == items_by_name["notes.md"]["fileId"]

    report_record = service.store.get_file(str(items_by_name["report.docx"]["fileId"]))
    metrics_record = service.store.get_file(str(items_by_name["metrics.xlsx"]["fileId"]))
    deck_record = service.store.get_file(str(items_by_name["deck.pptx"]["fileId"]))
    image_record = service.store.get_file(str(items_by_name["topology.png"]["fileId"]))
    notes_record = service.store.get_file(str(items_by_name["notes.md"]["fileId"]))
    assert report_record is not None
    assert metrics_record is not None
    assert deck_record is not None
    assert image_record is not None
    assert notes_record is not None

    assert report_record.processing_params["indexMode"] == "raganything_file"
    assert report_record.processing_params["chunkManifestCount"] >= 1
    assert metrics_record.processing_params["indexMode"] == "raganything_file"
    assert metrics_record.processing_params["chunkManifestCount"] >= 1
    assert deck_record.processing_params["indexMode"] == "raganything_file"
    assert image_record.processing_params["indexMode"] == "raganything_file"
    assert notes_record.processing_params["indexMode"] == "text_insert"
