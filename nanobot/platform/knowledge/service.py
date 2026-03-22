"""Service layer for the rebuilt knowledge-base subsystem."""

from __future__ import annotations

import asyncio
import csv
import hashlib
import io
import json
import math
import mimetypes
import re
import shutil
import textwrap
import uuid
from concurrent.futures import Future, ThreadPoolExecutor
from dataclasses import replace
from pathlib import Path
from threading import Lock
from typing import TYPE_CHECKING, Any
from urllib.parse import urlparse

import chardet
import httpx
from loguru import logger
from lxml import html as lxml_html
from openpyxl import load_workbook
from readability import Document as ReadabilityDocument

from nanobot.platform.instances import PlatformInstance
from nanobot.platform.knowledge.milvus_backend import MilvusBackendError, MilvusKnowledgeBackend
from nanobot.platform.knowledge.models import (
    KnowledgeBaseDefinition,
    KnowledgeDocumentStatus,
    KnowledgeFile,
    KnowledgeIngestJob,
    KnowledgeJobStatus,
    KnowledgeQueryParams,
    default_query_params_for_kb_type,
    default_query_params_payload,
    now_iso,
)
from nanobot.platform.knowledge.store import KnowledgeBaseStore
from nanobot.utils.helpers import ensure_dir, safe_filename

if TYPE_CHECKING:
    from nanobot.config.schema import Config
    from nanobot.platform.knowledge.rag_engine import RAGEngine


class KnowledgeBaseNotFoundError(KeyError):
    """Raised when a knowledge base does not exist."""


class KnowledgeBaseConflictError(RuntimeError):
    """Raised when a knowledge base name would conflict."""


class KnowledgeBaseValidationError(ValueError):
    """Raised when the payload or source data is invalid."""


class KnowledgeSourceNotFoundError(KeyError):
    """Raised when a knowledge file or folder does not exist."""


def _slugify(value: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return normalized or "knowledge-base"


def _short_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


class KnowledgeBaseService:
    """Instance-scoped knowledge-base CRUD, file tree, retrieval, and agent bindings."""

    def __init__(
        self,
        store: KnowledgeBaseStore,
        *,
        instance: PlatformInstance | None = None,
        instance_id: str = "default",
        tenant_id: str = "default",
        rag_engine: RAGEngine | None = None,
        max_background_jobs: int = 4,
        config: Config | None = None,
    ) -> None:
        self.store = store
        self.instance = instance
        self.instance_id = instance_id
        self.tenant_id = tenant_id
        self.rag_engine = rag_engine
        self.config = config
        self._executor = ThreadPoolExecutor(
            max_workers=max(1, max_background_jobs),
            thread_name_prefix=f"knowledge-{instance_id}",
        )
        self._futures: set[Future[Any]] = set()
        self._futures_lock = Lock()
        self._rag_engine_lock = Lock()
        self._vector_lock = Lock()
        self._evaluation_lock = Lock()
        self._job_options_lock = Lock()
        self._milvus_lock = Lock()
        self._job_options: dict[str, dict[str, Any]] = {}
        self._milvus_clients: dict[str, Any] = {}
        self._milvus_backend = MilvusKnowledgeBackend(self._kb_milvus_uri)

        import threading

        self._loop = asyncio.new_event_loop()
        self._loop_thread = threading.Thread(
            target=self._execute_loop,
            name=f"knowledge-loop-{instance_id}",
            daemon=True,
        )
        self._loop_thread.start()

    def _execute_loop(self) -> None:
        asyncio.set_event_loop(self._loop)
        self._loop.run_forever()

    def _run_async(self, coro: Any) -> Any:
        future = asyncio.run_coroutine_threadsafe(coro, self._loop)
        return future.result()

    def shutdown(self) -> None:
        if self.rag_engine is not None and self._loop.is_running():
            try:
                self._run_async(self.rag_engine.shutdown_async())
            except Exception:
                logger.exception("Failed to shut down knowledge RAG engine")
        try:
            self._milvus_backend.close_all()
        except Exception:
            logger.exception("Failed to close Milvus clients")
        if self._loop.is_running():
            self._loop.call_soon_threadsafe(self._loop.stop)
        if self._loop_thread.is_alive():
            self._loop_thread.join(timeout=1.0)
        self._executor.shutdown(wait=False, cancel_futures=True)

    def set_rag_engine(self, rag_engine: RAGEngine | None) -> None:
        with self._rag_engine_lock:
            self.rag_engine = rag_engine

    def set_config(self, config: Config | None) -> None:
        self.config = config

    def _track_future(self, future: Future[Any]) -> None:
        with self._futures_lock:
            self._futures.add(future)

        def _cleanup(done: Future[Any]) -> None:
            with self._futures_lock:
                self._futures.discard(done)
            try:
                done.result()
            except Exception:
                logger.exception("Knowledge background job crashed")

        future.add_done_callback(_cleanup)

    def _submit_background_job(self, fn: Any, *args: Any) -> None:
        future = self._executor.submit(fn, *args)
        self._track_future(future)

    def _store_job_options(self, job_id: str, payload: dict[str, Any]) -> None:
        with self._job_options_lock:
            self._job_options[job_id] = dict(payload)

    def _consume_job_options(self, job_id: str) -> dict[str, Any]:
        with self._job_options_lock:
            return dict(self._job_options.pop(job_id, {}))

    @staticmethod
    def _get_value(payload: dict[str, Any], *keys: str) -> Any:
        for key in keys:
            if key in payload:
                return payload[key]
        return None

    @staticmethod
    def _normalize_text(value: Any, *, required: bool = False, field_name: str = "value") -> str:
        text = str(value or "").strip()
        if required and not text:
            raise KnowledgeBaseValidationError(f"{field_name} is required.")
        return text

    @staticmethod
    def _normalize_string_list(value: Any, *, field_name: str) -> list[str]:
        if value is None:
            return []
        if not isinstance(value, list):
            raise KnowledgeBaseValidationError(f"{field_name} must be a list of strings.")
        result: list[str] = []
        seen: set[str] = set()
        for item in value:
            text = str(item or "").strip()
            if not text or text in seen:
                continue
            seen.add(text)
            result.append(text)
        return result

    @staticmethod
    def _normalize_object(value: Any, *, field_name: str) -> dict[str, Any]:
        if value is None:
            return {}
        if not isinstance(value, dict):
            raise KnowledgeBaseValidationError(f"{field_name} must be an object.")
        return dict(value)

    @staticmethod
    def _normalize_kb_type(value: Any) -> str:
        kb_type = str(value or "lightrag").strip().lower() or "lightrag"
        if kb_type not in {"lightrag", "milvus"}:
            raise KnowledgeBaseValidationError(f"Unsupported knowledge base type: {kb_type}.")
        return kb_type

    def _next_kb_id(self, name: str) -> str:
        base = _slugify(name)
        candidate = base
        counter = 2
        while self.store.get_kb(candidate) is not None:
            candidate = f"{base}-{counter}"
            counter += 1
        return candidate

    def _ensure_unique_name(self, name: str, *, exclude_kb_id: str | None = None) -> None:
        existing = self.store.get_kb_by_name(name, tenant_id=self.tenant_id, instance_id=self.instance_id)
        if existing is None:
            return
        if exclude_kb_id and existing.kb_id == exclude_kb_id:
            return
        raise KnowledgeBaseConflictError(f"Knowledge base name '{name}' already exists.")

    def _kb_raw_dir(self, kb_id: str) -> Path:
        if self.instance is not None:
            return ensure_dir(self.instance.knowledge_files_dir() / kb_id)
        raise KnowledgeBaseValidationError("Platform instance is required for knowledge files.")

    def _kb_parsed_dir(self, kb_id: str) -> Path:
        if self.instance is not None:
            return ensure_dir(self.instance.knowledge_parsed_dir() / kb_id)
        raise KnowledgeBaseValidationError("Platform instance is required for parsed knowledge files.")

    def _kb_vector_dir(self, kb_id: str) -> Path:
        if self.instance is not None:
            return ensure_dir(self.instance.runtime_dir("knowledge-vectors") / kb_id)
        raise KnowledgeBaseValidationError("Platform instance is required for vector knowledge files.")

    def _kb_eval_dir(self, kb_id: str) -> Path:
        if self.instance is not None:
            return ensure_dir(self.instance.runtime_dir("knowledge-evaluation") / kb_id)
        raise KnowledgeBaseValidationError("Platform instance is required for evaluation files.")

    def _kb_vector_index_path(self, kb_id: str) -> Path:
        return self._kb_vector_dir(kb_id) / "chunks.json"

    def _kb_milvus_uri(self, kb_id: str) -> str:
        return str(self._kb_vector_dir(kb_id) / "milvus.db")

    def _kb_benchmarks_dir(self, kb_id: str) -> Path:
        return ensure_dir(self._kb_eval_dir(kb_id) / "benchmarks")

    def _kb_results_dir(self, kb_id: str) -> Path:
        return ensure_dir(self._kb_eval_dir(kb_id) / "results")

    def _file_storage_paths(self, kb_id: str, file_id: str, filename: str) -> tuple[Path, Path]:
        safe_name = safe_filename(filename or f"{file_id}.txt")
        raw_path = self._kb_raw_dir(kb_id) / f"{file_id}-{safe_name}"
        parsed_path = self._kb_parsed_dir(kb_id) / f"{file_id}.md"
        return raw_path, parsed_path

    def require_kb(self, kb_id: str) -> KnowledgeBaseDefinition:
        kb = self.store.get_kb(kb_id)
        if kb is None or kb.instance_id != self.instance_id or kb.tenant_id != self.tenant_id:
            raise KnowledgeBaseNotFoundError(kb_id)
        return kb

    def _serialize_kb(self, kb: KnowledgeBaseDefinition) -> dict[str, Any]:
        payload = kb.to_dict()
        payload["stats"] = self.store.get_kb_stats(kb.kb_id)
        return payload

    def _serialize_file(self, file: KnowledgeFile) -> dict[str, Any]:
        return file.to_dict()

    def _serialize_job(self, job: KnowledgeIngestJob) -> dict[str, Any]:
        return job.to_dict()

    def _ensure_lightrag(self, kb: KnowledgeBaseDefinition, *, feature: str) -> None:
        if str(kb.kb_type or "lightrag").strip().lower() != "lightrag":
            raise KnowledgeBaseValidationError(f"{feature} currently only supports LightRAG knowledge bases.")
        if self.rag_engine is None:
            raise KnowledgeBaseValidationError(f"{feature} is unavailable because the RAG engine is not configured.")

    @staticmethod
    def _kb_type(kb: KnowledgeBaseDefinition) -> str:
        return str(kb.kb_type or "lightrag").strip().lower()

    def _ensure_file_ops_supported(self, kb: KnowledgeBaseDefinition, *, feature: str) -> None:
        if self._kb_type(kb) not in {"lightrag", "milvus"}:
            raise KnowledgeBaseValidationError(f"{feature} is unavailable for knowledge base type '{kb.kb_type}'.")

    def _ensure_evaluation_supported(self, kb: KnowledgeBaseDefinition) -> None:
        if self._kb_type(kb) != "milvus":
            raise KnowledgeBaseValidationError("Evaluation currently only supports Milvus knowledge bases.")

    @staticmethod
    def _atomic_write_json(path: Path, payload: Any) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp_path = path.with_suffix(f"{path.suffix}.tmp")
        tmp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp_path.replace(path)

    @staticmethod
    def _read_json(path: Path, default: Any) -> Any:
        if not path.exists():
            return default
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            logger.warning("Failed to decode JSON file {}, falling back to default", path)
            return default

    def _list_descendant_ids(self, kb_id: str, root_file_id: str) -> list[str]:
        files = self.store.list_files(kb_id)
        by_parent: dict[str | None, list[KnowledgeFile]] = {}
        for item in files:
            by_parent.setdefault(item.parent_id, []).append(item)
        result: list[str] = []
        stack = [root_file_id]
        while stack:
            current = stack.pop()
            result.append(current)
            for child in by_parent.get(current, []):
                stack.append(child.file_id)
        return result

    def _require_file(self, kb_id: str, file_id: str) -> KnowledgeFile:
        file = self.store.get_file(file_id)
        if file is None or file.kb_id != kb_id:
            raise KnowledgeSourceNotFoundError(file_id)
        return file

    def _ensure_parent_folder(self, kb_id: str, parent_id: str | None) -> KnowledgeFile | None:
        if not parent_id:
            return None
        parent = self._require_file(kb_id, parent_id)
        if not parent.is_folder:
            raise KnowledgeBaseValidationError("parentId must reference a folder.")
        return parent

    def _logical_path(self, parent: KnowledgeFile | None, filename: str) -> str:
        clean_name = filename.strip()
        if not clean_name:
            raise KnowledgeBaseValidationError("filename is required.")
        if parent is None:
            return f"/{clean_name}"
        return f"{parent.path.rstrip('/')}/{clean_name}"

    def _dedupe_filename(
        self,
        kb_id: str,
        parent_id: str | None,
        filename: str,
        *,
        exclude_file_id: str | None = None,
    ) -> str:
        files = self.store.list_files(kb_id)
        sibling_names = {
            item.filename.lower()
            for item in files
            if item.parent_id == parent_id and item.file_id != exclude_file_id
        }
        if filename.lower() not in sibling_names:
            return filename
        stem = Path(filename).stem
        suffix = Path(filename).suffix
        counter = 2
        while True:
            candidate = f"{stem}-{counter}{suffix}"
            if candidate.lower() not in sibling_names:
                return candidate
            counter += 1

    def _update_file(self, file: KnowledgeFile) -> KnowledgeFile:
        updated = self.store.update_file(file)
        if updated is None:
            raise KnowledgeSourceNotFoundError(file.file_id)
        return updated

    def _update_file_path_recursive(self, kb_id: str, file: KnowledgeFile, *, parent: KnowledgeFile | None) -> None:
        next_path = self._logical_path(parent, file.filename)
        updated = self._update_file(replace(file, path=next_path, updated_at=now_iso()))
        children = [item for item in self.store.list_files(kb_id) if item.parent_id == file.file_id]
        for child in children:
            self._update_file_path_recursive(kb_id, child, parent=updated)

    def _create_job(self, kb_id: str, job_kind: str, target_file_ids: list[str]) -> KnowledgeIngestJob:
        now = now_iso()
        return self.store.insert_job(
            KnowledgeIngestJob(
                job_id=_short_id("job"),
                tenant_id=self.tenant_id,
                instance_id=self.instance_id,
                kb_id=kb_id,
                job_kind=job_kind,
                target_file_ids=list(target_file_ids),
                status=KnowledgeJobStatus.QUEUED,
                track_id=_short_id("track"),
                created_at=now,
                updated_at=now,
            )
        )

    def _start_job(self, job: KnowledgeIngestJob) -> KnowledgeIngestJob:
        updated = replace(job, status=KnowledgeJobStatus.RUNNING, updated_at=now_iso())
        persisted = self.store.update_job(updated)
        if persisted is None:
            raise RuntimeError(f"Failed to start knowledge job {job.job_id}")
        return persisted

    def _finish_job(self, job: KnowledgeIngestJob, *, error_summary: str | None = None) -> KnowledgeIngestJob:
        updated = replace(
            job,
            status=KnowledgeJobStatus.FAILED if error_summary else KnowledgeJobStatus.SUCCEEDED,
            error_summary=error_summary,
            updated_at=now_iso(),
        )
        persisted = self.store.update_job(updated)
        if persisted is None:
            raise RuntimeError(f"Failed to finish knowledge job {job.job_id}")
        return persisted

    def list_knowledge_bases(self, enabled: bool | None = None) -> list[dict[str, Any]]:
        return [self._serialize_kb(item) for item in self.store.list_kbs(
            tenant_id=self.tenant_id,
            instance_id=self.instance_id,
            enabled=enabled,
        )]

    def list_accessible_knowledge_bases(self, enabled: bool | None = True) -> list[dict[str, Any]]:
        return self.list_knowledge_bases(enabled=enabled)

    def get_knowledge_base(self, kb_id: str) -> dict[str, Any]:
        return self._serialize_kb(self.require_kb(kb_id))

    def create_knowledge_base(self, payload: dict[str, Any]) -> dict[str, Any]:
        name = self._normalize_text(self._get_value(payload, "name"), required=True, field_name="name")
        self._ensure_unique_name(name)
        now = now_iso()
        kb_type = self._normalize_kb_type(self._get_value(payload, "kbType", "kb_type"))
        query_payload = self._get_value(payload, "queryParams", "query_params", "retrievalProfile", "retrieval_profile")
        kb = self.store.create_kb(
            KnowledgeBaseDefinition(
                kb_id=self._next_kb_id(name),
                tenant_id=self.tenant_id,
                instance_id=self.instance_id,
                name=name,
                description=self._normalize_text(self._get_value(payload, "description"), field_name="description"),
                enabled=True if self._get_value(payload, "enabled") is None else bool(self._get_value(payload, "enabled")),
                kb_type=kb_type,
                embed_info=self._normalize_object(
                    self._get_value(payload, "embedInfo", "embed_info") or {},
                    field_name="embedInfo",
                ),
                llm_info=self._normalize_object(
                    self._get_value(payload, "llmInfo", "llm_info") or {},
                    field_name="llmInfo",
                ),
                query_params=KnowledgeQueryParams.from_dict(
                    query_payload,
                    defaults=default_query_params_payload(kb_type),
                ),
                additional_params=self._normalize_object(
                    self._get_value(payload, "additionalParams", "additional_params") or {},
                    field_name="additionalParams",
                ),
                share_config=self._normalize_object(
                    self._get_value(payload, "shareConfig", "share_config") or {},
                    field_name="shareConfig",
                ),
                sample_questions=self._normalize_string_list(
                    self._get_value(payload, "sampleQuestions", "sample_questions"),
                    field_name="sampleQuestions",
                ),
                tags=self._normalize_string_list(self._get_value(payload, "tags"), field_name="tags"),
                created_at=now,
                updated_at=now,
            )
        )
        return self._serialize_kb(kb)

    def update_knowledge_base(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        existing = self.require_kb(kb_id)
        name = self._normalize_text(
            self._get_value(payload, "name") if "name" in payload else existing.name,
            required=True,
            field_name="name",
        )
        self._ensure_unique_name(name, exclude_kb_id=kb_id)
        next_kb_type = (
            self._normalize_kb_type(self._get_value(payload, "kbType", "kb_type"))
            if ("kbType" in payload or "kb_type" in payload)
            else existing.kb_type
        )
        has_query_params_update = (
            "queryParams" in payload
            or "query_params" in payload
            or "retrievalProfile" in payload
            or "retrieval_profile" in payload
        )
        updated = replace(
            existing,
            name=name,
            description=(
                self._normalize_text(self._get_value(payload, "description"), field_name="description")
                if "description" in payload
                else existing.description
            ),
            enabled=bool(self._get_value(payload, "enabled")) if "enabled" in payload else existing.enabled,
            kb_type=next_kb_type,
            embed_info=(
                self._normalize_object(self._get_value(payload, "embedInfo", "embed_info"), field_name="embedInfo")
            ) if ("embedInfo" in payload or "embed_info" in payload) else existing.embed_info,
            llm_info=(
                self._normalize_object(self._get_value(payload, "llmInfo", "llm_info"), field_name="llmInfo")
            ) if ("llmInfo" in payload or "llm_info" in payload) else existing.llm_info,
            query_params=(
                KnowledgeQueryParams.from_dict(
                    self._get_value(payload, "queryParams", "query_params", "retrievalProfile", "retrieval_profile"),
                    defaults=default_query_params_payload(next_kb_type),
                )
            ) if has_query_params_update else (
                default_query_params_for_kb_type(next_kb_type)
                if str(next_kb_type).strip().lower() != self._kb_type(existing)
                else existing.query_params
            ),
            additional_params=(
                self._normalize_object(self._get_value(payload, "additionalParams", "additional_params"), field_name="additionalParams")
            ) if ("additionalParams" in payload or "additional_params" in payload) else existing.additional_params,
            share_config=(
                self._normalize_object(self._get_value(payload, "shareConfig", "share_config"), field_name="shareConfig")
            ) if ("shareConfig" in payload or "share_config" in payload) else existing.share_config,
            mindmap=self._get_value(payload, "mindmap") if "mindmap" in payload else existing.mindmap,
            sample_questions=(
                self._normalize_string_list(self._get_value(payload, "sampleQuestions", "sample_questions"), field_name="sampleQuestions")
            ) if ("sampleQuestions" in payload or "sample_questions" in payload) else existing.sample_questions,
            tags=(
                self._normalize_string_list(self._get_value(payload, "tags"), field_name="tags")
            ) if "tags" in payload else existing.tags,
            updated_at=now_iso(),
        )
        persisted = self.store.update_kb(updated)
        if persisted is None:
            raise KnowledgeBaseNotFoundError(kb_id)
        return self._serialize_kb(persisted)

    def delete_knowledge_base(self, kb_id: str) -> bool:
        kb = self.require_kb(kb_id)
        if self._kb_type(kb) == "lightrag" and self.rag_engine is not None:
            try:
                self._run_async(self.rag_engine.delete_kb(kb_id))
            except Exception:
                logger.exception("Failed to delete RAG data for knowledge base {}", kb_id)
        try:
            self._milvus_backend.close_kb(kb_id)
        except Exception:
            logger.exception("Failed to close Milvus client for {}", kb_id)
        for root in (
            self._kb_raw_dir(kb_id),
            self._kb_parsed_dir(kb_id),
            self._kb_vector_dir(kb_id),
            self._kb_eval_dir(kb_id),
        ):
            if root.exists():
                shutil.rmtree(root, ignore_errors=True)
        return self.store.delete_kb(kb_id)

    def resolve_bound_kbs(self, kb_ids: list[str]) -> list[KnowledgeBaseDefinition]:
        result: list[KnowledgeBaseDefinition] = []
        missing: list[str] = []
        for kb_id in kb_ids:
            kb = self.store.get_kb(kb_id)
            if kb is None or kb.instance_id != self.instance_id or kb.tenant_id != self.tenant_id or not kb.enabled:
                missing.append(kb_id)
                continue
            result.append(kb)
        if missing:
            raise KnowledgeBaseValidationError(
                f"Agent references unknown or disabled knowledge bases: {', '.join(missing)}"
            )
        return result

    def list_files(self, kb_id: str) -> dict[str, Any]:
        self.require_kb(kb_id)
        files = self.store.list_files(kb_id)
        return {
            "items": [self._serialize_file(item) for item in files],
            "stats": self.store.get_kb_stats(kb_id),
        }

    def create_folder(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="Folder creation")
        parent = self._ensure_parent_folder(kb_id, self._normalize_text(payload.get("parentId"), field_name="parentId") or None)
        name = self._normalize_text(self._get_value(payload, "name", "filename"), required=True, field_name="name")
        filename = self._dedupe_filename(kb_id, parent.file_id if parent else None, name)
        now = now_iso()
        folder = self.store.insert_file(
            KnowledgeFile(
                file_id=_short_id("file"),
                kb_id=kb_id,
                tenant_id=self.tenant_id,
                instance_id=self.instance_id,
                parent_id=parent.file_id if parent else None,
                filename=filename,
                original_filename=filename,
                file_type="folder",
                path=self._logical_path(parent, filename),
                raw_path=None,
                markdown_file=None,
                status=KnowledgeDocumentStatus.FOLDER,
                is_folder=True,
                created_at=now,
                updated_at=now,
            )
        )
        return self._serialize_file(folder)

    def upload_files(self, kb_id: str, files: list[dict[str, Any]], *, parent_id: str | None = None) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="File upload")
        parent = self._ensure_parent_folder(kb_id, parent_id)
        created: list[dict[str, Any]] = []
        for item in files:
            content = item.get("content")
            if not isinstance(content, (bytes, bytearray)) or not content:
                raise KnowledgeBaseValidationError("Uploaded knowledge file content is required.")
            file_name = self._normalize_text(item.get("file_name"), required=True, field_name="file_name")
            mime_type = self._normalize_text(item.get("mime_type"), field_name="mime_type") or None
            filename = self._dedupe_filename(kb_id, parent.file_id if parent else None, file_name)
            file_id = _short_id("file")
            raw_path, parsed_path = self._file_storage_paths(kb_id, file_id, filename)
            raw_path.write_bytes(bytes(content))
            now = now_iso()
            record = self.store.insert_file(
                KnowledgeFile(
                    file_id=file_id,
                    kb_id=kb_id,
                    tenant_id=self.tenant_id,
                    instance_id=self.instance_id,
                    parent_id=parent.file_id if parent else None,
                    filename=filename,
                    original_filename=file_name,
                    file_type=(Path(filename).suffix.lower().lstrip(".") or "file"),
                    path=self._logical_path(parent, filename),
                    raw_path=str(raw_path),
                    markdown_file=str(parsed_path),
                    status=KnowledgeDocumentStatus.UPLOADED,
                    content_hash=hashlib.sha256(bytes(content)).hexdigest(),
                    file_size=len(content),
                    content_type=mime_type,
                    processing_params={"sourceType": "upload"},
                    is_folder=False,
                    created_at=now,
                    updated_at=now,
                )
            )
            created.append(self._serialize_file(record))
        return {"items": created}

    def fetch_url_file(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="URL import")
        url = self._normalize_text(self._get_value(payload, "url"), required=True, field_name="url")
        parent = self._ensure_parent_folder(kb_id, self._normalize_text(payload.get("parentId"), field_name="parentId") or None)

        response = httpx.get(url, timeout=20.0, follow_redirects=True)
        response.raise_for_status()
        content_type = str(response.headers.get("content-type") or "").split(";")[0].strip() or None
        suffix = Path(urlparse(url).path).suffix
        if not suffix and content_type:
            suffix = mimetypes.guess_extension(content_type) or ".html"
        base_name = Path(urlparse(url).path).name or "web-source"
        filename = base_name if Path(base_name).suffix else f"{base_name}{suffix or '.html'}"
        filename = self._dedupe_filename(kb_id, parent.file_id if parent else None, safe_filename(filename))
        file_id = _short_id("file")
        raw_path, parsed_path = self._file_storage_paths(kb_id, file_id, filename)
        raw_path.write_bytes(response.content)
        now = now_iso()
        record = self.store.insert_file(
            KnowledgeFile(
                file_id=file_id,
                kb_id=kb_id,
                tenant_id=self.tenant_id,
                instance_id=self.instance_id,
                parent_id=parent.file_id if parent else None,
                filename=filename,
                original_filename=filename,
                file_type="web_url",
                path=self._logical_path(parent, filename),
                raw_path=str(raw_path),
                markdown_file=str(parsed_path),
                status=KnowledgeDocumentStatus.UPLOADED,
                content_hash=hashlib.sha256(response.content).hexdigest(),
                file_size=len(response.content),
                content_type=content_type,
                processing_params={"sourceType": "web_url", "sourceUrl": url},
                created_at=now,
                updated_at=now,
            )
        )
        return self._serialize_file(record)

    def add_source_file(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        source_type = self._normalize_text(payload.get("sourceType"), required=True, field_name="sourceType")
        if source_type == "web_url":
            return self.fetch_url_file(kb_id, payload)
        if source_type != "faq_table":
            raise KnowledgeBaseValidationError(f"Unsupported knowledge source type: {source_type}")

        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="Knowledge source import")
        items = payload.get("items")
        if not isinstance(items, list) or not items:
            raise KnowledgeBaseValidationError("FAQ table source requires a non-empty items list.")
        normalized_items: list[dict[str, str]] = []
        for index, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                raise KnowledgeBaseValidationError(f"FAQ item {index} must be an object.")
            question = self._normalize_text(item.get("question"), required=True, field_name=f"items[{index}].question")
            answer = self._normalize_text(item.get("answer"), required=True, field_name=f"items[{index}].answer")
            normalized_items.append({"question": question, "answer": answer})

        parent = self._ensure_parent_folder(kb_id, self._normalize_text(payload.get("parentId"), field_name="parentId") or None)
        title = self._normalize_text(payload.get("title"), field_name="title") or "faq-table"
        filename = self._dedupe_filename(kb_id, parent.file_id if parent else None, safe_filename(f"{title}.json"))
        file_id = _short_id("file")
        raw_path, parsed_path = self._file_storage_paths(kb_id, file_id, filename)
        raw_bytes = json.dumps(normalized_items, ensure_ascii=False, indent=2).encode("utf-8")
        raw_path.write_bytes(raw_bytes)
        now = now_iso()
        record = self.store.insert_file(
            KnowledgeFile(
                file_id=file_id,
                kb_id=kb_id,
                tenant_id=self.tenant_id,
                instance_id=self.instance_id,
                parent_id=parent.file_id if parent else None,
                filename=filename,
                original_filename=filename,
                file_type="faq_table",
                path=self._logical_path(parent, filename),
                raw_path=str(raw_path),
                markdown_file=str(parsed_path),
                status=KnowledgeDocumentStatus.UPLOADED,
                content_hash=hashlib.sha256(raw_bytes).hexdigest(),
                file_size=len(raw_bytes),
                content_type="application/json",
                processing_params={"sourceType": "faq_table", "faqItems": normalized_items},
                created_at=now,
                updated_at=now,
            )
        )
        return self._serialize_file(record)

    def move_file(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="File move")
        file_id = self._normalize_text(payload.get("fileId"), required=True, field_name="fileId")
        target_parent_id = self._normalize_text(payload.get("targetParentId"), field_name="targetParentId") or None
        rename_to = self._normalize_text(payload.get("filename"), field_name="filename") or None

        file = self._require_file(kb_id, file_id)
        parent = self._ensure_parent_folder(kb_id, target_parent_id)
        if file.is_folder and parent is not None:
            descendants = set(self._list_descendant_ids(kb_id, file.file_id))
            if parent.file_id in descendants:
                raise KnowledgeBaseValidationError("Cannot move a folder into its own descendant.")

        filename = self._dedupe_filename(
            kb_id,
            parent.file_id if parent else None,
            rename_to or file.filename,
            exclude_file_id=file.file_id,
        )
        updated = self._update_file(
            replace(
                file,
                parent_id=parent.file_id if parent else None,
                filename=filename,
                path=self._logical_path(parent, filename),
                updated_at=now_iso(),
            )
        )
        for child in [item for item in self.store.list_files(kb_id) if item.parent_id == updated.file_id]:
            self._update_file_path_recursive(kb_id, child, parent=updated)
        return self._serialize_file(updated)

    def delete_file(self, kb_id: str, file_id: str) -> bool:
        kb = self.require_kb(kb_id)
        descendants = list(reversed(self._list_descendant_ids(kb_id, file_id)))
        for target_id in descendants:
            file = self._require_file(kb_id, target_id)
            if (
                not file.is_folder
                and file.status == KnowledgeDocumentStatus.INDEXED
                and self._kb_type(kb) == "lightrag"
                and self.rag_engine is not None
            ):
                try:
                    self._run_async(self.rag_engine.delete_document(kb_id, file.file_id))
                except Exception:
                    logger.exception("Failed to delete indexed knowledge file {}", file.file_id)
            if not file.is_folder and self._kb_type(kb) == "milvus":
                self._remove_vector_entries_for_file(kb_id, file.file_id)
            for candidate in (file.raw_path, file.markdown_file):
                if candidate:
                    path = Path(candidate)
                    if path.exists():
                        try:
                            path.unlink()
                        except OSError:
                            logger.warning("Failed to remove knowledge artifact {}", candidate)
            self.store.delete_file(target_id)
        return True

    def delete_files(self, kb_id: str, file_ids: list[str]) -> dict[str, Any]:
        deleted: list[str] = []
        for file_id in self._normalize_string_list(file_ids, field_name="fileIds"):
            if self.delete_file(kb_id, file_id):
                deleted.append(file_id)
        return {"deletedCount": len(deleted), "fileIds": deleted}

    def list_jobs(self, kb_id: str) -> list[dict[str, Any]]:
        self.require_kb(kb_id)
        return [self._serialize_job(job) for job in self.store.list_jobs(kb_id)]

    def _select_target_files(self, kb_id: str, file_ids: list[str] | None = None) -> list[KnowledgeFile]:
        self.require_kb(kb_id)
        files = self.store.list_files(kb_id)
        if file_ids:
            wanted = set(self._normalize_string_list(file_ids, field_name="fileIds"))
            selected = [item for item in files if item.file_id in wanted and not item.is_folder]
        else:
            selected = [item for item in files if not item.is_folder]
        if not selected:
            raise KnowledgeBaseValidationError("No knowledge files matched the requested operation.")
        return selected

    def _resolve_existing_files(
        self,
        kb_id: str,
        candidates: list[str],
    ) -> tuple[list[KnowledgeFile], list[str]]:
        self.require_kb(kb_id)
        files = self.store.list_files(kb_id)
        resolved: list[KnowledgeFile] = []
        seen: set[str] = set()
        missing: list[str] = []
        for raw_candidate in candidates:
            candidate = str(raw_candidate or "").strip()
            if not candidate:
                continue
            matched: KnowledgeFile | None = None
            for item in files:
                tokens = {
                    item.file_id,
                    item.filename,
                    item.path,
                    item.raw_path or "",
                    Path(item.raw_path).name if item.raw_path else "",
                }
                if candidate in tokens:
                    matched = item
                    break
            if matched is None:
                missing.append(candidate)
                continue
            if matched.file_id in seen or matched.is_folder:
                continue
            seen.add(matched.file_id)
            resolved.append(matched)
        return resolved, missing

    def _extract_requested_file_ids(self, payload: dict[str, Any] | None) -> list[str] | None:
        if not isinstance(payload, dict):
            return None
        for key in ("fileIds", "file_ids", "docIds", "doc_ids"):
            value = payload.get(key)
            if isinstance(value, list):
                return self._normalize_string_list(value, field_name=key)
        return None

    @staticmethod
    def _merge_query_payload(payload: dict[str, Any] | None) -> dict[str, Any]:
        merged = dict(payload or {})
        meta = merged.pop("meta", None)
        if isinstance(meta, dict):
            return {
                **meta,
                **merged,
            }
        return merged

    @staticmethod
    def _detect_encoding(content: bytes) -> str:
        detection = chardet.detect(content)
        return str(detection.get("encoding") or "utf-8")

    def _decode_text(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode(self._detect_encoding(content), errors="ignore")

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"[ \t]+\n", "\n", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _html_to_text(self, raw_html: str) -> tuple[str, str | None]:
        doc = ReadabilityDocument(raw_html)
        title = doc.short_title() or None
        summary_html = doc.summary(html_partial=True)
        text = lxml_html.fromstring(summary_html).text_content()
        return self._normalize_whitespace(text), title

    def _json_to_text(self, raw_json: str) -> tuple[str, list[dict[str, Any]] | None]:
        payload = json.loads(raw_json)
        if isinstance(payload, list) and payload and all(isinstance(item, dict) for item in payload):
            faq_items = []
            for item in payload:
                question = str(item.get("question") or item.get("q") or "").strip()
                answer = str(item.get("answer") or item.get("a") or "").strip()
                if question and answer:
                    faq_items.append({"question": question, "answer": answer})
            if faq_items:
                return "\n\n".join(f"Q: {item['question']}\nA: {item['answer']}" for item in faq_items), faq_items
        return json.dumps(payload, ensure_ascii=False, indent=2), None

    def _csv_to_text(self, raw_csv: str) -> tuple[str, list[dict[str, Any]] | None]:
        reader = csv.DictReader(io.StringIO(raw_csv))
        rows = list(reader)
        faq_items = []
        lines: list[str] = []
        for row in rows:
            question = str(row.get("question") or row.get("q") or "").strip()
            answer = str(row.get("answer") or row.get("a") or "").strip()
            if question and answer:
                faq_items.append({"question": question, "answer": answer})
            if row:
                lines.append(" | ".join(f"{key}: {value}" for key, value in row.items()))
        if faq_items:
            return "\n\n".join(f"Q: {item['question']}\nA: {item['answer']}" for item in faq_items), faq_items
        return "\n".join(lines), None

    def _xlsx_to_text(self, content: bytes) -> str:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(cell or "").strip() for cell in rows[0]]
            sections.append(f"# Sheet: {sheet.title}")
            for row in rows[1:]:
                line = " | ".join(
                    f"{header[index] or f'column_{index + 1}'}: {str(value or '').strip()}"
                    for index, value in enumerate(row)
                    if str(value or "").strip()
                )
                if line:
                    sections.append(line)
        return "\n".join(sections).strip()

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeBaseValidationError("PDF parsing requires optional dependency 'pypdf'.") from exc
        reader = PdfReader(io.BytesIO(content))
        return self._normalize_whitespace("\n\n".join(page.extract_text() or "" for page in reader.pages))

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise KnowledgeBaseValidationError("DOCX parsing requires optional dependency 'python-docx'.") from exc
        document = Document(io.BytesIO(content))
        return self._normalize_whitespace("\n".join(paragraph.text for paragraph in document.paragraphs))

    def _parse_file_content(
        self,
        *,
        title: str,
        file_name: str,
        mime_type: str | None,
        content: bytes,
    ) -> tuple[str, str, dict[str, Any], list[dict[str, Any]] | None]:
        suffix = Path(file_name).suffix.lower()
        parser_name = suffix.lstrip(".") or "text"
        metadata: dict[str, Any] = {}
        faq_items: list[dict[str, Any]] | None = None

        if suffix in {".txt", ".md"}:
            text = self._decode_text(content)
        elif suffix in {".html", ".htm"}:
            text, detected_title = self._html_to_text(self._decode_text(content))
            if detected_title and detected_title != title:
                metadata["detectedTitle"] = detected_title
        elif suffix == ".json":
            text, faq_items = self._json_to_text(self._decode_text(content))
        elif suffix == ".csv":
            text, faq_items = self._csv_to_text(self._decode_text(content))
        elif suffix == ".xlsx":
            text = self._xlsx_to_text(content)
        elif suffix == ".pdf":
            text = self._extract_pdf_text(content)
        elif suffix == ".docx":
            text = self._extract_docx_text(content)
        else:
            if mime_type and mime_type.startswith("text/"):
                text = self._decode_text(content)
            else:
                raise KnowledgeBaseValidationError(
                    f"Unsupported knowledge file type: {suffix or mime_type or file_name}"
                )

        normalized = self._normalize_whitespace(text)
        if not normalized:
            raise KnowledgeBaseValidationError("Parsed knowledge document is empty.")
        return normalized, parser_name, metadata, faq_items

    def _parse_single_file(self, file: KnowledgeFile) -> KnowledgeFile:
        current = self._require_file(file.kb_id, file.file_id)
        if current.is_folder:
            return current
        if not current.raw_path:
            raise KnowledgeBaseValidationError(f"Knowledge file {current.file_id} has no source file.")
        source_path = Path(current.raw_path)
        if not source_path.exists():
            raise KnowledgeBaseValidationError(f"Knowledge source file is missing: {source_path}")

        current = self._update_file(
            replace(current, status=KnowledgeDocumentStatus.PARSING, error_message=None, updated_at=now_iso())
        )
        content = source_path.read_bytes()
        text, parser_name, metadata, faq_items = self._parse_file_content(
            title=current.filename,
            file_name=current.original_filename or current.filename,
            mime_type=current.content_type,
            content=content,
        )
        parsed_path = Path(current.markdown_file or self._file_storage_paths(current.kb_id, current.file_id, current.filename)[1])
        parsed_path.parent.mkdir(parents=True, exist_ok=True)
        parsed_path.write_text(text, encoding="utf-8")

        processing_params = dict(current.processing_params)
        processing_params.update(metadata)
        processing_params["parserName"] = parser_name
        processing_params["parsedAt"] = now_iso()
        if faq_items:
            processing_params["faqItems"] = faq_items
        return self._update_file(
            replace(
                current,
                markdown_file=str(parsed_path),
                status=KnowledgeDocumentStatus.PARSED,
                processing_params=processing_params,
                error_message=None,
                updated_at=now_iso(),
            )
        )

    def _index_single_file(self, file: KnowledgeFile) -> KnowledgeFile:
        current = self._require_file(file.kb_id, file.file_id)
        kb = self.require_kb(current.kb_id)
        kb_type = self._kb_type(kb)
        if current.is_folder:
            return current
        if current.status not in {KnowledgeDocumentStatus.PARSED, KnowledgeDocumentStatus.INDEXED}:
            raise KnowledgeBaseValidationError(
                f"Knowledge file {current.file_id} must be parsed before indexing."
            )
        if not current.markdown_file:
            raise KnowledgeBaseValidationError(f"Knowledge file {current.file_id} has not been parsed.")

        parsed_path = Path(current.markdown_file)
        if not parsed_path.exists():
            raise KnowledgeBaseValidationError(f"Parsed knowledge file is missing: {parsed_path}")

        current = self._update_file(
            replace(current, status=KnowledgeDocumentStatus.INDEXING, error_message=None, updated_at=now_iso())
        )
        text = parsed_path.read_text(encoding="utf-8")

        if kb_type == "milvus":
            entries = self._build_milvus_chunk_entries(kb, current, text)
            self._milvus_backend.delete_file_entries(current.kb_id, current.file_id)
            self._milvus_backend.upsert_entries(current.kb_id, entries)
            processing_params = dict(current.processing_params)
            processing_params["chunksCount"] = len(entries)
            processing_params["indexedAt"] = now_iso()
            processing_params["indexBackend"] = "pymilvus-lite"
            return self._update_file(
                replace(
                    current,
                    status=KnowledgeDocumentStatus.INDEXED,
                    processing_params=processing_params,
                    error_message=None,
                    updated_at=now_iso(),
                )
            )

        self._ensure_lightrag(kb, feature="Knowledge indexing")
        if current.processing_params.get("chunksCount"):
            try:
                self._run_async(self.rag_engine.delete_document(current.kb_id, current.file_id))
            except Exception:
                logger.warning("Failed to delete old index for {}", current.file_id)
        chunk_texts = self._build_chunk_texts(kb, current, text)
        insert_chunks = getattr(self.rag_engine, "insert_chunks", None)
        if callable(insert_chunks):
            result = self._run_async(
                insert_chunks(
                    current.kb_id,
                    chunk_texts,
                    doc_id=current.file_id,
                    file_path=current.raw_path or current.filename,
                )
            )
        else:
            result = self._run_async(
                self.rag_engine.insert_text(
                    current.kb_id,
                    "\n\n".join(chunk_texts),
                    doc_id=current.file_id,
                    file_path=current.raw_path or current.filename,
                )
            )
        if not result.success:
            raise KnowledgeBaseValidationError(result.error or f"Failed to index knowledge file {current.file_id}")

        processing_params = dict(current.processing_params)
        processing_params["chunksCount"] = int(result.metadata.get("chunks_count") or len(chunk_texts))
        processing_params["indexedAt"] = now_iso()
        return self._update_file(
            replace(
                current,
                status=KnowledgeDocumentStatus.INDEXED,
                processing_params=processing_params,
                error_message=None,
                updated_at=now_iso(),
            )
        )

    def _run_parse_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                self._parse_single_file(self._require_file(job.kb_id, file_id))
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_PARSING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update parse error state for {}", file_id)
        self._finish_job(job, error_summary="; ".join(errors) if errors else None)

    def _run_index_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                self._index_single_file(self._require_file(job.kb_id, file_id))
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_INDEXING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update index error state for {}", file_id)
        finished_job = self._finish_job(job, error_summary="; ".join(errors) if errors else None)
        if finished_job.status == KnowledgeJobStatus.SUCCEEDED:
            try:
                kb = self.require_kb(job.kb_id)
                if bool((kb.additional_params or {}).get("auto_generate_questions")):
                    self.generate_sample_questions(job.kb_id, count=10)
            except Exception:
                logger.exception("Failed to auto-generate sample questions for {}", job.kb_id)

    def _run_ingest_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        options = self._consume_job_options(job_id)
        auto_index = bool(options.get("auto_index"))
        index_params = self._normalize_index_params(options.get("index_params"))
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                parsed = self._parse_single_file(self._require_file(job.kb_id, file_id))
                if not auto_index:
                    continue
                indexed_input = parsed
                if index_params:
                    indexed_input = self._update_file(
                        replace(
                            parsed,
                            processing_params={
                                **(parsed.processing_params or {}),
                                **index_params,
                            },
                            updated_at=now_iso(),
                        )
                    )
                self._index_single_file(indexed_input)
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_INDEXING if auto_index else KnowledgeDocumentStatus.ERROR_PARSING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update ingest error state for {}", file_id)
        finished_job = self._finish_job(job, error_summary="; ".join(errors) if errors else None)
        if auto_index and finished_job.status == KnowledgeJobStatus.SUCCEEDED:
            try:
                kb = self.require_kb(job.kb_id)
                if bool((kb.additional_params or {}).get("auto_generate_questions")):
                    self.generate_sample_questions(job.kb_id, count=10)
            except Exception:
                logger.exception("Failed to auto-generate sample questions for {}", job.kb_id)

    def parse_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="Knowledge parsing")
        requested_file_ids = self._extract_requested_file_ids(payload)
        if not requested_file_ids:
            raise KnowledgeBaseValidationError("fileIds is required for parse operations.")
        selected = self._select_target_files(kb_id, requested_file_ids)
        job = self._create_job(kb_id, "parse", [item.file_id for item in selected])
        self._submit_background_job(self._run_parse_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }

    def index_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="Knowledge indexing")
        requested_file_ids = self._extract_requested_file_ids(payload)
        if not requested_file_ids:
            raise KnowledgeBaseValidationError("fileIds is required for index operations.")
        selected = self._select_target_files(kb_id, requested_file_ids)
        index_params = self._normalize_index_params(payload.get("params"))
        if index_params:
            refreshed_selected: list[KnowledgeFile] = []
            for item in selected:
                if item.is_folder:
                    refreshed_selected.append(item)
                    continue
                persisted = self._update_file(
                    replace(
                        item,
                        processing_params={
                            **(item.processing_params or {}),
                            **index_params,
                        },
                        updated_at=now_iso(),
                    )
                )
                refreshed_selected.append(persisted)
            selected = refreshed_selected
        job = self._create_job(kb_id, "index", [item.file_id for item in selected])
        self._submit_background_job(self._run_index_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }

    def ingest_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_file_ops_supported(kb, feature="Knowledge ingest")
        params = self._normalize_object(payload.get("params"), field_name="params")
        requested_file_ids = self._extract_requested_file_ids(payload)
        if requested_file_ids:
            selected = self._select_target_files(kb_id, requested_file_ids)
        else:
            items = self._normalize_string_list(payload.get("items"), field_name="items")
            if not items:
                raise KnowledgeBaseValidationError("items is required for ingest operations.")
            selected, missing = self._resolve_existing_files(kb_id, items)
            if missing:
                raise KnowledgeBaseValidationError(f"Unknown uploaded items: {', '.join(missing)}.")
            if not selected:
                raise KnowledgeBaseValidationError("No uploaded items matched the requested ingest payload.")

        parent_id = (
            self._normalize_text(params.get("parentId"), field_name="parentId")
            or self._normalize_text(params.get("parent_id"), field_name="parent_id")
            or None
        )
        if parent_id:
            refreshed: list[KnowledgeFile] = []
            for item in selected:
                moved = self.move_file(
                    kb_id,
                    {
                        "fileId": item.file_id,
                        "parentId": parent_id,
                    },
                )
                refreshed.append(self._require_file(kb_id, str(moved.get("fileId") or item.file_id)))
            selected = refreshed

        job = self._create_job(kb_id, "ingest", [item.file_id for item in selected])
        self._store_job_options(
            job.job_id,
            {
                "auto_index": bool(params.get("auto_index") or params.get("autoIndex")),
                "index_params": params,
            },
        )
        self._submit_background_job(self._run_ingest_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }

    def get_query_params(self, kb_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        return kb.query_params.to_dict()

    def get_query_param_schema(self, kb_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        if self._kb_type(kb) == "milvus":
            return {
                "type": "milvus",
                "options": [
                    {
                        "key": "mode",
                        "label": "检索模式",
                        "type": "select",
                        "default": "vector",
                        "options": [
                            {"value": "vector", "label": "向量检索"},
                            {"value": "keyword", "label": "关键词检索"},
                            {"value": "hybrid", "label": "混合检索"},
                        ],
                    },
                    {
                        "key": "topK",
                        "label": "最终返回 Chunk 数",
                        "type": "number",
                        "default": 10,
                        "min": 1,
                        "max": 100,
                    },
                    {
                        "key": "similarity_threshold",
                        "label": "相似度阈值",
                        "type": "number",
                        "default": 0.0,
                        "min": 0.0,
                        "max": 1.0,
                        "step": 0.05,
                    },
                    {
                        "key": "keyword_top_k",
                        "label": "关键词候选数",
                        "type": "number",
                        "default": 50,
                        "min": 1,
                        "max": 200,
                    },
                ],
            }
        return {
            "type": "lightrag",
            "options": [
                {
                    "key": "mode",
                    "label": "检索模式",
                    "type": "select",
                    "default": "mix",
                    "options": [
                        {"value": "local", "label": "Local"},
                        {"value": "global", "label": "Global"},
                        {"value": "hybrid", "label": "Hybrid"},
                        {"value": "naive", "label": "Naive"},
                        {"value": "mix", "label": "Mix"},
                    ],
                },
                {
                    "key": "topK",
                    "label": "TopK",
                    "type": "number",
                    "default": 10,
                    "min": 1,
                    "max": 100,
                },
                {
                    "key": "chunkTopK",
                    "label": "Chunk TopK",
                    "type": "number",
                    "default": 12,
                    "min": 1,
                    "max": 100,
                },
                {
                    "key": "responseType",
                    "label": "回答格式",
                    "type": "select",
                    "default": "Multiple Paragraphs",
                    "options": [
                        {"value": "Single Paragraph", "label": "Single Paragraph"},
                        {"value": "Multiple Paragraphs", "label": "Multiple Paragraphs"},
                        {"value": "Bullet Points", "label": "Bullet Points"},
                    ],
                },
                {
                    "key": "onlyNeedContext",
                    "label": "只返回上下文",
                    "type": "boolean",
                    "default": True,
                },
                {
                    "key": "onlyNeedPrompt",
                    "label": "只返回提示词",
                    "type": "boolean",
                    "default": False,
                },
                {
                    "key": "enableRerank",
                    "label": "启用重排",
                    "type": "boolean",
                    "default": False,
                },
            ],
        }

    def update_query_params(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        existing = kb.query_params.to_dict()
        merged_options = dict(kb.query_params.options)
        merged_payload: dict[str, Any] = {**existing, "options": merged_options}

        for key, value in dict(payload or {}).items():
            if key == "options" and isinstance(value, dict):
                merged_options.update(value)
                continue
            if value is None:
                continue
            if key in {
                "mode",
                "topK",
                "top_k",
                "chunkTopK",
                "chunk_top_k",
                "responseType",
                "response_type",
                "onlyNeedContext",
                "only_need_context",
                "onlyNeedPrompt",
                "only_need_prompt",
                "enableRerank",
                "enable_rerank",
                "rerankModel",
                "rerank_model",
            }:
                merged_payload[key] = value
                continue
            merged_options[key] = value

        updated = self.store.update_kb(
            replace(
                kb,
                query_params=KnowledgeQueryParams.from_dict(
                    merged_payload,
                    defaults=default_query_params_payload(self._kb_type(kb)),
                ),
                updated_at=now_iso(),
            )
        )
        if updated is None:
            raise KnowledgeBaseNotFoundError(kb_id)
        return updated.query_params.to_dict()

    def generate_description(self, payload: dict[str, Any]) -> dict[str, Any]:
        name = self._normalize_text(payload.get("name"), required=True, field_name="name")
        current_description = self._normalize_text(
            payload.get("currentDescription") if "currentDescription" in payload else payload.get("current_description"),
            field_name="currentDescription",
        )
        file_list = self._normalize_string_list(
            payload.get("fileList") if "fileList" in payload else payload.get("file_list"),
            field_name="fileList",
        )
        kb_id = self._normalize_text(payload.get("kbId") if "kbId" in payload else payload.get("kb_id"), field_name="kbId")
        if kb_id and not file_list:
            kb = self.require_kb(kb_id)
            file_list = [
                item.path
                for item in self.store.list_files(kb.kb_id)
                if not item.is_folder and str(item.path or "").strip()
            ]
        display_files = file_list[:50]
        file_text = "\n".join(f"- {item}" for item in display_files)
        more_text = f"\n... (还有 {len(file_list) - len(display_files)} 个文件)" if len(file_list) > len(display_files) else ""
        file_rule = "5. 请结合提供的文件列表来概括知识范围。" if display_files else ""
        file_section = f"文件列表:\n{file_text}{more_text}" if display_files else ""

        prompt = textwrap.dedent(
            f"""
            请帮我优化以下知识库的描述。

            知识库名称: {name}
            当前描述: {current_description or "暂无描述"}

            要求:
            1. 这个描述将作为智能体工具的描述使用。
            2. 要清晰说明知识库包含什么内容、适合解答什么问题。
            3. 描述保持简洁有力，通常 2 到 4 句话。
            4. 不要使用 Markdown，不要添加前缀说明。
            {file_rule}

            {file_section}
            """
        ).strip()
        description = self._generate_with_llm(
            system_prompt="你是知识库信息架构专家，擅长为 AI Agent 编写清晰、准确、可检索的知识库描述。",
            user_prompt=prompt,
        )
        if description:
            return {"description": description}

        display_names = [Path(item).name or item for item in display_files[:6]]
        if display_names:
            summary = f"{name}知识库，主要包含{'、'.join(display_names)}等资料。适合用于回答与这些资料相关的操作说明、事实查询和流程问题。"
        elif current_description:
            summary = current_description
        else:
            summary = f"{name}知识库，适合用于回答与{name}相关的文档检索、流程说明和事实问答。"
        return {"description": summary}

    def _normalize_index_params(self, payload: Any) -> dict[str, Any]:
        if not isinstance(payload, dict):
            return {}
        normalized: dict[str, Any] = {}
        raw_chunk_size = payload.get("chunk_size") if "chunk_size" in payload else payload.get("chunkSize")
        if raw_chunk_size is not None:
            normalized["chunk_size"] = max(200, int(raw_chunk_size))
        raw_chunk_overlap = payload.get("chunk_overlap") if "chunk_overlap" in payload else payload.get("chunkOverlap")
        if raw_chunk_overlap is not None:
            overlap = max(0, int(raw_chunk_overlap))
            chunk_size = int(normalized.get("chunk_size") or payload.get("chunk_size") or payload.get("chunkSize") or 1000)
            normalized["chunk_overlap"] = min(overlap, max(0, chunk_size - 1))
        chunk_preset = self._normalize_text(
            payload.get("chunk_preset_id") if "chunk_preset_id" in payload else payload.get("chunkPresetId"),
            field_name="chunkPresetId",
        )
        if chunk_preset:
            normalized["chunk_preset_id"] = chunk_preset
        qa_separator = self._normalize_text(
            payload.get("qa_separator") if "qa_separator" in payload else payload.get("qaSeparator"),
            field_name="qaSeparator",
        )
        if qa_separator:
            normalized["qa_separator"] = qa_separator
        return normalized

    def _matching_file_tokens(self, kb_id: str, file_ids: list[str] | None = None, file_name: str | None = None) -> set[str]:
        tokens: set[str] = set()
        files = self.store.list_files(kb_id)
        for file in files:
            if file_ids and file.file_id not in set(file_ids):
                continue
            if file_name and file_name.lower() not in file.filename.lower():
                continue
            tokens.add(file.file_id)
            tokens.add(file.filename)
            if file.raw_path:
                tokens.add(file.raw_path)
                tokens.add(Path(file.raw_path).name)
        return {token for token in tokens if token}

    @staticmethod
    def _filter_query_result(raw: dict[str, Any], *, tokens: set[str]) -> dict[str, Any]:
        if not tokens:
            return raw
        result = json.loads(json.dumps(raw, ensure_ascii=False))
        data = result.get("data") or {}
        references = list(data.get("references") or [])
        filtered_references = [
            item for item in references
            if any(token in str(item.get("file_path") or "") or token == str(item.get("reference_id") or "") for token in tokens)
        ]
        reference_ids = {str(item.get("reference_id") or "") for item in filtered_references}
        filtered_chunks = [
            item for item in (data.get("chunks") or [])
            if str(item.get("reference_id") or "") in reference_ids
            or any(token in str(item.get("file_path") or "") for token in tokens)
        ]
        filtered_entities = [
            item for item in (data.get("entities") or [])
            if any(token in str(item.get("file_path") or "") for token in tokens)
            or any(token in str(item.get("source_id") or "") for token in tokens)
        ]
        filtered_relationships = [
            item for item in (data.get("relationships") or [])
            if any(token in str(item.get("file_path") or "") for token in tokens)
            or any(token in str(item.get("source_id") or "") for token in tokens)
        ]
        data["references"] = filtered_references
        data["chunks"] = filtered_chunks
        data["entities"] = filtered_entities
        data["relationships"] = filtered_relationships
        result["data"] = data
        result.setdefault("metadata", {})
        result["metadata"]["fileFilterApplied"] = True
        result["metadata"]["filteredReferenceCount"] = len(filtered_references)
        return result

    def _query_milvus_database(
        self,
        kb: KnowledgeBaseDefinition,
        query_text: str,
        params: KnowledgeQueryParams,
        *,
        file_ids: list[str] | None = None,
        file_name: str | None = None,
    ) -> dict[str, Any]:
        all_entries = self._load_vector_entries(kb.kb_id)
        entries = list(all_entries)
        if file_ids:
            wanted = set(file_ids)
            entries = [item for item in entries if str(item.get("fileId") or "") in wanted]
        if file_name:
            lowered = file_name.lower()
            entries = [
                item
                for item in entries
                if lowered in str(item.get("filename") or "").lower()
                or lowered in str(item.get("path") or "").lower()
            ]
        filtered_chunk_ids = [str(item.get("chunkId") or "") for item in entries if str(item.get("chunkId") or "").strip()]

        options = dict(params.options or {})
        search_mode = str(options.get("search_mode") or options.get("searchMode") or params.mode or "vector").lower()
        if search_mode not in {"vector", "keyword", "hybrid"}:
            search_mode = "vector"
        final_top_k = max(1, int(options.get("final_top_k") or options.get("finalTopK") or params.top_k or 8))
        similarity_threshold = float(
            options.get("similarity_threshold") or options.get("similarityThreshold") or 0.0
        )
        keyword_top_k = max(1, int(options.get("keyword_top_k") or options.get("keywordTopK") or max(final_top_k, 20)))
        recall_top_k = max(final_top_k, int(options.get("recall_top_k") or options.get("recallTopK") or final_top_k))

        vector_results: list[dict[str, Any]] = []
        if search_mode in {"vector", "hybrid"} and entries:
            query_embedding = self._embed_texts([query_text], kb=kb)[0]
            scoped_chunk_ids = None
            if filtered_chunk_ids and len(filtered_chunk_ids) != len(all_entries):
                scoped_chunk_ids = filtered_chunk_ids
            for item in self._milvus_backend.search_entries(
                kb.kb_id,
                query_embedding,
                chunk_ids=scoped_chunk_ids,
                limit=recall_top_k,
            ):
                similarity = float(item.get("similarity") or item.get("score") or 0.0)
                if similarity < similarity_threshold:
                    continue
                vector_results.append(dict(item))

        keyword_results: list[dict[str, Any]] = []
        if search_mode in {"keyword", "hybrid"} and entries:
            query_tokens = set(self._keyword_tokens(query_text))
            if query_tokens:
                scored: list[tuple[dict[str, Any], int]] = []
                for item in entries:
                    item_tokens = set(item.get("tokens") or [])
                    overlap = len(query_tokens & item_tokens)
                    if overlap > 0:
                        scored.append((item, overlap))
                if scored:
                    max_overlap = max(score for _, score in scored)
                    for item, overlap in scored:
                        keyword_score = overlap / max_overlap if max_overlap else 0.0
                        keyword_results.append(
                            {
                                **item,
                                "score": keyword_score,
                                "keyword_score": keyword_score,
                            }
                        )
                    keyword_results.sort(key=lambda item: float(item.get("score") or 0.0), reverse=True)
                    keyword_results = keyword_results[:keyword_top_k]

        if search_mode == "vector":
            retrieved = vector_results
        elif search_mode == "keyword":
            retrieved = keyword_results
        else:
            merged: dict[str, dict[str, Any]] = {}
            for item in vector_results:
                merged[str(item.get("chunkId") or uuid.uuid4().hex)] = dict(item)
            for item in keyword_results:
                key = str(item.get("chunkId") or uuid.uuid4().hex)
                if key in merged:
                    merged[key]["keyword_score"] = item.get("keyword_score")
                    merged[key]["score"] = max(
                        float(merged[key].get("score") or 0.0),
                        float(item.get("keyword_score") or item.get("score") or 0.0),
                    )
                else:
                    merged[key] = dict(item)
            retrieved = sorted(merged.values(), key=lambda item: float(item.get("score") or 0.0), reverse=True)

        references_map: dict[str, dict[str, Any]] = {}
        chunk_payloads: list[dict[str, Any]] = []
        for item in retrieved[:final_top_k]:
            file_id = str(item.get("fileId") or "")
            reference_id = file_id or str(item.get("chunkId") or "")
            if reference_id and reference_id not in references_map:
                references_map[reference_id] = {
                    "reference_id": reference_id,
                    "file_path": item.get("filename") or item.get("filePath") or item.get("path"),
                    "file_name": item.get("filename"),
                }
            chunk_payloads.append(
                {
                    "chunk_id": item.get("chunkId"),
                    "chunkId": item.get("chunkId"),
                    "content": item.get("content"),
                    "file_id": file_id,
                    "fileId": file_id,
                    "filename": item.get("filename"),
                    "reference_id": reference_id,
                    "file_path": item.get("filename") or item.get("filePath") or item.get("path"),
                    "chunk_index": item.get("chunkIndex"),
                    "chunkIndex": item.get("chunkIndex"),
                    "score": item.get("score"),
                    "keyword_score": item.get("keyword_score"),
                    "similarity": item.get("similarity"),
                    "metadata": {
                        "chunk_id": item.get("chunkId"),
                        "source": item.get("filename") or item.get("filePath") or item.get("path"),
                        "file_id": file_id,
                    },
                }
            )

        message = None
        if not params.only_need_context and not params.only_need_prompt:
            message = self._synthesize_answer_from_chunks(kb, query_text, chunk_payloads)

        return {
            "status": "success",
            "message": message,
            "query": query_text,
            "data": {
                "entities": [],
                "relationships": [],
                "chunks": chunk_payloads,
                "references": list(references_map.values()),
            },
            "metadata": {
                "mode": search_mode,
                "query_mode": search_mode,
                "kbType": "milvus",
                "candidateCount": len(retrieved),
                "fileFilterApplied": bool(file_ids or file_name),
            },
            "queryParams": params.to_dict(),
        }

    def query_database(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        merged_payload = self._merge_query_payload(payload)
        query_text = self._normalize_text(self._get_value(merged_payload, "query", "queryText"), required=True, field_name="query")
        file_ids = self._extract_requested_file_ids(merged_payload) or []
        file_name = self._normalize_text(self._get_value(merged_payload, "fileName", "file_name"), field_name="fileName") or None
        params = KnowledgeQueryParams.from_dict(
            {
                **kb.query_params.to_dict(),
                **merged_payload,
            },
            defaults=default_query_params_payload(self._kb_type(kb)),
        )

        if self._kb_type(kb) == "milvus":
            return self._query_milvus_database(
                kb,
                query_text,
                params,
                file_ids=file_ids or None,
                file_name=file_name,
            )

        self._ensure_lightrag(kb, feature="Knowledge query")
        raw = self._run_async(
            self.rag_engine.query_structured(
                kb_id,
                query_text,
                mode=params.mode,
                top_k=params.top_k,
                chunk_top_k=params.chunk_top_k,
                response_type=params.response_type,
                only_need_context=params.only_need_context,
                only_need_prompt=params.only_need_prompt,
                enable_rerank=params.enable_rerank,
            )
        )
        tokens = self._matching_file_tokens(kb_id, file_ids=file_ids or None, file_name=file_name)
        filtered = self._filter_query_result(raw, tokens=tokens)
        filtered["query"] = query_text
        filtered["queryParams"] = params.to_dict()
        return filtered

    def retrieve(
        self,
        kb_ids: list[str],
        query: str,
        *,
        limit: int | None = None,
        filters: dict[str, Any] | None = None,
        requested_mode: str | None = None,
    ) -> dict[str, Any]:
        resolved = self.resolve_bound_kbs(kb_ids)
        if not query.strip():
            raise KnowledgeBaseValidationError("query is required.")
        files_by_kb = {
            item.kb_id: self.store.list_files(item.kb_id)
            for item in resolved
        }

        def _match_file(kb_id: str, file_path: str | None) -> KnowledgeFile | None:
            candidate = str(file_path or "")
            if not candidate:
                return None
            for item in files_by_kb.get(kb_id, []):
                if item.raw_path and candidate in {item.raw_path, Path(item.raw_path).name}:
                    return item
                if item.file_id in candidate or item.filename in candidate:
                    return item
            return None

        items: list[dict[str, Any]] = []
        effective_modes: list[str] = []
        for kb in resolved:
            kb_type = self._kb_type(kb)
            resolved_mode = str(requested_mode or "").strip()
            if not resolved_mode:
                if kb_type == "milvus":
                    resolved_mode = str(
                        kb.query_params.options.get("search_mode")
                        or kb.query_params.options.get("searchMode")
                        or kb.query_params.mode
                        or "vector"
                    ).strip() or "vector"
                else:
                    resolved_mode = str(kb.query_params.mode or "mix").strip() or "mix"
            effective_modes.append(resolved_mode)

            if kb_type == "lightrag":
                if self.rag_engine is None:
                    continue
                hits = self._run_async(
                    self.rag_engine.query(
                        [kb.kb_id],
                        query,
                        mode=resolved_mode,
                        top_k=max(1, int(limit or 8)),
                    )
                )
                for hit in hits:
                    file = _match_file(kb.kb_id, str(hit.metadata.get("file_path") or ""))
                    items.append(
                        {
                            "kbId": kb.kb_id,
                            "kbName": kb.name,
                            "docId": file.file_id if file is not None else None,
                            "title": file.filename if file is not None else kb.name,
                            "content": hit.content,
                            "preview": hit.content[:280],
                            "score": hit.score,
                            "metadata": dict(hit.metadata),
                            "citation": {
                                "kbId": kb.kb_id,
                                "kbName": kb.name,
                                "docId": file.file_id if file is not None else None,
                                "title": file.filename if file is not None else kb.name,
                                "sourceType": file.file_type if file is not None else "knowledge",
                                "sourceUri": (file.processing_params or {}).get("sourceUrl") if file is not None else None,
                                "fileName": file.filename if file is not None else None,
                                "mimeType": file.content_type if file is not None else None,
                                "chunkOrdinal": None,
                            },
                        }
                    )
                continue

            query_result = self._query_milvus_database(
                kb,
                query,
                KnowledgeQueryParams.from_dict(
                    {
                        **kb.query_params.to_dict(),
                        "mode": resolved_mode,
                        "topK": max(1, int(limit or 8)),
                    },
                    defaults=default_query_params_payload(kb_type),
                ),
            )
            for chunk in list((query_result.get("data") or {}).get("chunks") or [])[: max(1, int(limit or 8))]:
                file = _match_file(kb.kb_id, str(chunk.get("file_path") or ""))
                content = str(chunk.get("content") or "")
                items.append(
                    {
                        "kbId": kb.kb_id,
                        "kbName": kb.name,
                        "docId": file.file_id if file is not None else None,
                        "title": file.filename if file is not None else kb.name,
                        "content": content,
                        "preview": content[:280],
                        "score": float(chunk.get("score") or 0.0),
                        "metadata": {
                            "kb_id": kb.kb_id,
                            "file_path": chunk.get("file_path"),
                            "chunk_id": chunk.get("chunk_id"),
                        },
                        "citation": {
                            "kbId": kb.kb_id,
                            "kbName": kb.name,
                            "docId": file.file_id if file is not None else None,
                            "title": file.filename if file is not None else kb.name,
                            "sourceType": file.file_type if file is not None else "knowledge",
                            "sourceUri": (file.processing_params or {}).get("sourceUrl") if file is not None else None,
                            "fileName": file.filename if file is not None else None,
                            "mimeType": file.content_type if file is not None else None,
                            "chunkOrdinal": None,
                        },
                    }
                )

        items.sort(key=lambda item: float(item.get("score") or 0.0), reverse=True)
        if limit:
            items = items[: max(1, int(limit))]

        return {
            "hits": items,
            "requestedMode": requested_mode or "auto",
            "effectiveMode": (
                effective_modes[0]
                if effective_modes and len(set(effective_modes)) == 1
                else (requested_mode or "mixed")
            ),
            "filters": dict(filters or {}),
        }

    @staticmethod
    def _extract_headings(text: str, *, limit: int = 6) -> list[str]:
        headings: list[str] = []
        for line in text.splitlines():
            stripped = line.strip().lstrip("#").strip()
            if not stripped:
                continue
            if len(stripped) > 100:
                stripped = stripped[:100].rstrip()
            if stripped not in headings:
                headings.append(stripped)
            if len(headings) >= limit:
                break
        return headings

    def _load_outline(self, file: KnowledgeFile) -> dict[str, Any]:
        text = ""
        if file.markdown_file and Path(file.markdown_file).exists():
            try:
                text = Path(file.markdown_file).read_text(encoding="utf-8")
            except OSError:
                text = ""
        headings = self._extract_headings(text, limit=5)
        return {
            "fileId": file.file_id,
            "filename": file.filename,
            "path": file.path,
            "headings": headings,
            "excerpt": text[:1200].strip(),
        }

    def _provider_from_config(self):
        if self.config is None:
            return None
        from nanobot.providers.base import GenerationSettings
        from nanobot.providers.custom_provider import CustomProvider
        from nanobot.providers.litellm_provider import LiteLLMProvider
        from nanobot.providers.openai_codex_provider import OpenAICodexProvider

        model = self.config.agents.defaults.model
        provider_name = self.config.get_provider_name(model)
        provider_cfg = self.config.get_provider(model)

        if provider_name == "openai_codex" or model.startswith("openai-codex/"):
            provider = OpenAICodexProvider(default_model=model)
        elif provider_name == "custom":
            provider = CustomProvider(
                api_key=(provider_cfg.api_key if provider_cfg and provider_cfg.api_key else "no-key"),
                api_base=self.config.get_api_base(model) or "http://localhost:8000/v1",
                default_model=model,
            )
        elif provider_name == "azure_openai":
            if provider_cfg and provider_cfg.api_key and provider_cfg.api_base:
                from nanobot.providers.azure_openai_provider import AzureOpenAIProvider

                provider = AzureOpenAIProvider(
                    api_key=provider_cfg.api_key,
                    api_base=provider_cfg.api_base,
                    default_model=model,
                )
            else:
                provider = LiteLLMProvider(
                    api_key=None,
                    api_base=None,
                    default_model=model,
                    provider_name=provider_name,
                )
        else:
            provider = LiteLLMProvider(
                api_key=provider_cfg.api_key if provider_cfg and provider_cfg.api_key else None,
                api_base=self.config.get_api_base(model),
                default_model=model,
                extra_headers=provider_cfg.extra_headers if provider_cfg else None,
                provider_name=provider_name,
            )

        defaults = self.config.agents.defaults
        provider.generation = GenerationSettings(
            temperature=0.2,
            max_tokens=min(defaults.max_tokens, 2000),
            reasoning_effort=defaults.reasoning_effort,
        )
        return provider

    @staticmethod
    def _extract_json_text(raw: str) -> str | None:
        text = str(raw or "").strip()
        if not text:
            return None
        if text.startswith("{") or text.startswith("["):
            return text
        match = re.search(r"```(?:json)?\s*([\s\S]+?)```", text)
        if match:
            return match.group(1).strip()
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return text[start:end + 1]
        return None

    def _generate_with_llm(self, *, system_prompt: str, user_prompt: str) -> str | None:
        try:
            provider = self._provider_from_config()
            if provider is None:
                return None
            response = self._run_async(
                provider.chat_with_retry(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    model=self.config.agents.defaults.model if self.config is not None else None,
                )
            )
            content = str(response.content or "").strip()
            return content or None
        except Exception:
            logger.exception("Knowledge LLM generation failed")
            return None

    async def _aembed_texts(
        self,
        texts: list[str],
        *,
        kb: KnowledgeBaseDefinition | None = None,
    ) -> list[list[float]]:
        if not texts:
            return []
        if self.config is None:
            raise KnowledgeBaseValidationError("Embedding is unavailable because the runtime config is missing.")

        from litellm import aembedding

        from nanobot.platform.knowledge.rag_engine import (
            RAGEngine,
            _first_binding_name_by_capability,
            _resolve_binding_runtime,
        )

        rag_config = getattr(self.config, "rag", None)
        binding_name = (
            str(getattr(rag_config, "embedding_binding", None) or "").strip()
            or _first_binding_name_by_capability(self.config, "embedding")
            or None
        )
        runtime = _resolve_binding_runtime(self.config, binding_name=binding_name, model=None)
        requested_model = (
            str((kb.embed_info or {}).get("modelName") or (kb.embed_info or {}).get("model") or "").strip()
            if kb is not None
            else ""
        )
        model = requested_model or str(getattr(runtime.get("binding"), "model", None) or "text-embedding-3-large")
        provider_name = runtime.get("provider_name")
        api_key = str(runtime.get("api_key") or "")
        api_base = str(runtime.get("api_base") or "") or None
        extra_headers = dict(runtime.get("extra_headers") or {})
        resolved_model, custom_llm_provider, resolved_api_base = RAGEngine._resolve_litellm_runtime(
            model=model,
            provider_name=provider_name,
            api_key=api_key,
            api_base=api_base,
            request_type="embedding",
        )

        vectors: list[list[float]] = []
        batch_size = 32
        for offset in range(0, len(texts), batch_size):
            batch = texts[offset : offset + batch_size]
            kw: dict[str, Any] = {"model": resolved_model, "input": batch}
            if api_key:
                kw["api_key"] = api_key
            if resolved_api_base:
                kw["api_base"] = resolved_api_base
            if custom_llm_provider:
                kw["custom_llm_provider"] = custom_llm_provider
            if provider_name == "dashscope" and custom_llm_provider == "openai":
                kw["encoding_format"] = "float"
            if extra_headers:
                kw["extra_headers"] = extra_headers
            response = await aembedding(**kw)
            vectors.extend([list(item["embedding"]) for item in response.data])
        return vectors

    def _embed_texts(
        self,
        texts: list[str],
        *,
        kb: KnowledgeBaseDefinition | None = None,
    ) -> list[list[float]]:
        return self._run_async(self._aembed_texts(texts, kb=kb))

    @staticmethod
    def _keyword_tokens(text: str) -> list[str]:
        tokens: set[str] = set()
        for raw in re.findall(r"[a-zA-Z0-9_]+|[\u4e00-\u9fff]+", text.lower()):
            raw = raw.strip()
            if not raw:
                continue
            if re.fullmatch(r"[\u4e00-\u9fff]+", raw):
                tokens.add(raw)
                if len(raw) > 2:
                    for index in range(len(raw) - 1):
                        tokens.add(raw[index : index + 2])
            elif len(raw) > 1:
                tokens.add(raw)
        return sorted(tokens)

    @staticmethod
    def _cosine_similarity(left: list[float], right: list[float]) -> float:
        if not left or not right or len(left) != len(right):
            return 0.0
        dot = sum(a * b for a, b in zip(left, right))
        left_norm = math.sqrt(sum(a * a for a in left))
        right_norm = math.sqrt(sum(b * b for b in right))
        if left_norm == 0 or right_norm == 0:
            return 0.0
        return dot / (left_norm * right_norm)

    @staticmethod
    def _milvus_collection_name() -> str:
        return "knowledge_chunks"

    def _ensure_milvus_runtime(self) -> None:
        try:
            self._milvus_backend.ensure_runtime()
        except MilvusBackendError as exc:
            raise KnowledgeBaseValidationError(str(exc)) from exc

    def _get_milvus_client(self, kb_id: str) -> Any:
        self._ensure_milvus_runtime()
        return self._milvus_backend.get_client(kb_id)

    def _ensure_milvus_collection(self, kb: KnowledgeBaseDefinition, *, dimension: int) -> Any:
        self._ensure_milvus_runtime()
        self._milvus_backend.ensure_collection(kb.kb_id, dimension=dimension)
        return self._milvus_backend.get_client(kb.kb_id)

    @staticmethod
    def _strip_embedding_entries(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
        stored: list[dict[str, Any]] = []
        for item in entries:
            if not isinstance(item, dict):
                continue
            payload = dict(item)
            payload.pop("embedding", None)
            stored.append(payload)
        return stored

    def _load_vector_entries(self, kb_id: str) -> list[dict[str, Any]]:
        return self._milvus_backend.list_entries(kb_id)

    def _save_vector_entries(self, kb_id: str, entries: list[dict[str, Any]]) -> None:
        self._milvus_backend.replace_entries(kb_id, entries)

    def _remove_vector_entries_for_file(self, kb_id: str, file_id: str) -> None:
        try:
            self._milvus_backend.delete_file_entries(kb_id, file_id)
        except Exception:
            logger.exception("Failed to remove Milvus entries for {}", file_id)

    @staticmethod
    def _split_large_block(text: str, *, chunk_size: int, chunk_overlap: int) -> list[str]:
        if len(text) <= chunk_size:
            return [text]
        result: list[str] = []
        start = 0
        step = max(1, chunk_size - max(0, chunk_overlap))
        while start < len(text):
            result.append(text[start : start + chunk_size].strip())
            start += step
        return [item for item in result if item]

    def _chunk_plain_text(self, text: str, *, chunk_size: int, chunk_overlap: int) -> list[str]:
        paragraphs = [item.strip() for item in re.split(r"\n{2,}", text) if item.strip()]
        if not paragraphs:
            paragraphs = [text.strip()]
        chunks: list[str] = []
        current = ""
        for paragraph in paragraphs:
            if len(paragraph) > chunk_size:
                if current.strip():
                    chunks.append(current.strip())
                    current = ""
                chunks.extend(self._split_large_block(paragraph, chunk_size=chunk_size, chunk_overlap=chunk_overlap))
                continue
            candidate = paragraph if not current else f"{current}\n\n{paragraph}"
            if len(candidate) <= chunk_size:
                current = candidate
                continue
            if current.strip():
                chunks.append(current.strip())
            if chunk_overlap > 0 and chunks:
                overlap = chunks[-1][-chunk_overlap:].strip()
                current = f"{overlap}\n\n{paragraph}".strip() if overlap else paragraph
            else:
                current = paragraph
        if current.strip():
            chunks.append(current.strip())
        return chunks

    @staticmethod
    def _chunk_by_headings(text: str) -> list[str]:
        sections: list[str] = []
        current: list[str] = []
        for line in text.splitlines():
            stripped = line.strip()
            if re.match(r"^(#{1,6}\s+|chapter\s+\d+|section\s+\d+|第.+[章节篇])", stripped, re.IGNORECASE):
                if current:
                    sections.append("\n".join(current).strip())
                    current = []
            if stripped:
                current.append(line)
        if current:
            sections.append("\n".join(current).strip())
        return [item for item in sections if item]

    @staticmethod
    def _chunk_qa_text(text: str) -> list[str]:
        normalized = text.replace("\r\n", "\n")
        pattern = re.compile(
            r"(?:^|\n)(?:Q[:：]\s*(?P<q>.+?)\nA[:：]\s*(?P<a>.+?))(?=\nQ[:：]|\Z)",
            re.DOTALL | re.IGNORECASE,
        )
        result = []
        for match in pattern.finditer(normalized):
            question = str(match.group("q") or "").strip()
            answer = str(match.group("a") or "").strip()
            if question and answer:
                result.append(f"Q: {question}\nA: {answer}")
        return result

    @staticmethod
    def _chunk_law_text(text: str) -> list[str]:
        parts = re.split(r"(?=第[\d一二三四五六七八九十百千]+条)", text)
        result = [item.strip() for item in parts if item.strip()]
        return result

    def _build_chunk_texts(
        self,
        kb: KnowledgeBaseDefinition,
        file: KnowledgeFile,
        text: str,
    ) -> list[str]:
        params = {**(kb.additional_params or {}), **(file.processing_params or {})}
        chunk_size = max(200, int(params.get("chunk_size") or params.get("chunkSize") or 1000))
        chunk_overlap = max(0, int(params.get("chunk_overlap") or params.get("chunkOverlap") or 200))
        qa_separator = str(params.get("qa_separator") or params.get("qaSeparator") or "").strip()
        chunk_preset_id = str(params.get("chunk_preset_id") or params.get("chunkPresetId") or "general").strip().lower()
        faq_items = file.processing_params.get("faqItems")

        chunk_texts: list[str] = []
        if isinstance(faq_items, list) and faq_items:
            for item in faq_items:
                if not isinstance(item, dict):
                    continue
                question = str(item.get("question") or "").strip()
                answer = str(item.get("answer") or "").strip()
                if question and answer:
                    chunk_texts.append(f"Q: {question}\nA: {answer}")
        elif chunk_preset_id == "qa":
            chunk_texts = self._chunk_qa_text(text)
        elif qa_separator and qa_separator in text:
            chunk_texts = [item.strip() for item in text.split(qa_separator) if item.strip()]
        elif chunk_preset_id == "book":
            chunk_texts = self._chunk_by_headings(text)
        elif chunk_preset_id == "laws":
            chunk_texts = self._chunk_law_text(text)
        else:
            chunk_texts = self._chunk_plain_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

        if not chunk_texts:
            raise KnowledgeBaseValidationError("No chunks were produced for the selected knowledge file.")
        return chunk_texts

    def _build_milvus_chunk_entries(
        self,
        kb: KnowledgeBaseDefinition,
        file: KnowledgeFile,
        text: str,
    ) -> list[dict[str, Any]]:
        chunk_texts = self._build_chunk_texts(kb, file, text)
        embeddings = self._embed_texts(chunk_texts, kb=kb)
        entries: list[dict[str, Any]] = []
        for index, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings), start=1):
            chunk_id = f"{file.file_id}::chunk::{index:04d}"
            entries.append(
                {
                    "chunkId": chunk_id,
                    "kbId": kb.kb_id,
                    "fileId": file.file_id,
                    "filename": file.filename,
                    "path": file.path,
                    "filePath": file.raw_path or file.filename,
                    "chunkIndex": index,
                    "content": chunk_text,
                    "embedding": embedding,
                    "tokens": self._keyword_tokens(chunk_text),
                    "createdAt": now_iso(),
                }
            )
        return entries

    def _synthesize_answer_from_chunks(
        self,
        kb: KnowledgeBaseDefinition,
        query_text: str,
        chunks: list[dict[str, Any]],
    ) -> str | None:
        if not chunks:
            return None
        excerpts = [
            {
                "chunkId": item.get("chunk_id") or item.get("chunkId"),
                "file": item.get("file_path") or item.get("filename") or item.get("filePath"),
                "content": str(item.get("content") or "").strip()[:1600],
            }
            for item in chunks[:6]
        ]
        llm_answer = self._generate_with_llm(
            system_prompt=(
                "你是知识库问答助手。只能依据提供的检索片段回答。"
                "如果片段不足以支持结论，要明确说信息不足。"
            ),
            user_prompt=json.dumps(
                {
                    "knowledgeBase": kb.name,
                    "query": query_text,
                    "chunks": excerpts,
                },
                ensure_ascii=False,
            ),
        )
        if llm_answer:
            return llm_answer
        fallback = "\n\n".join(str(item.get("content") or "").strip() for item in chunks[:2]).strip()
        return fallback[:1600] if fallback else None

    def _build_mindmap_fallback(self, kb: KnowledgeBaseDefinition, files: list[KnowledgeFile]) -> dict[str, Any]:
        by_parent: dict[str | None, list[KnowledgeFile]] = {}
        for item in files:
            by_parent.setdefault(item.parent_id, []).append(item)

        def _node_for(file: KnowledgeFile) -> dict[str, Any]:
            if file.is_folder:
                return {
                    "content": file.filename,
                    "children": [_node_for(child) for child in by_parent.get(file.file_id, [])],
                }
            outline = self._load_outline(file)
            children = [{"content": heading, "children": []} for heading in outline["headings"][:4]]
            return {"content": file.filename, "children": children}

        root_children = [_node_for(item) for item in by_parent.get(None, [])]
        return {"content": kb.name, "children": root_children}

    def generate_mindmap(self, kb_id: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        payload = payload or {}
        file_ids = self._normalize_string_list(payload.get("fileIds"), field_name="fileIds") if isinstance(payload.get("fileIds"), list) else []
        files = [item for item in self.store.list_files(kb_id) if not file_ids or item.file_id in set(file_ids)]
        fallback = self._build_mindmap_fallback(kb, files)
        outlines = [self._load_outline(item) for item in files if not item.is_folder][:12]
        llm_raw = self._generate_with_llm(
            system_prompt="你是知识架构整理助手。只返回 JSON，结构必须是 {\"content\": string, \"children\": []}。",
            user_prompt=json.dumps(
                {
                    "knowledgeBase": kb.name,
                    "description": kb.description,
                    "files": outlines,
                },
                ensure_ascii=False,
            ),
        )
        generated = fallback
        if llm_raw:
            json_text = self._extract_json_text(llm_raw)
            if json_text:
                try:
                    candidate = json.loads(json_text)
                    if isinstance(candidate, dict) and candidate.get("content"):
                        generated = candidate
                except json.JSONDecodeError:
                    logger.warning("Mindmap JSON decode failed, using fallback")

        updated = self.store.update_kb(replace(kb, mindmap=generated, updated_at=now_iso()))
        if updated is None:
            raise KnowledgeBaseNotFoundError(kb_id)
        return {"mindmap": generated}

    def get_mindmap(self, kb_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        if kb.mindmap is None:
            raise KnowledgeBaseValidationError("Knowledge mindmap has not been generated yet.")
        return {"mindmap": kb.mindmap}

    def _build_sample_question_fallback(self, kb: KnowledgeBaseDefinition, files: list[KnowledgeFile], count: int) -> list[str]:
        questions: list[str] = []
        for file in files:
            if file.is_folder:
                continue
            outline = self._load_outline(file)
            title = outline["filename"]
            if title:
                questions.append(f"知识库里《{title}》的核心内容是什么？")
                questions.append(f"如果我需要用到《{title}》，最关键的步骤有哪些？")
            for heading in outline["headings"][:3]:
                questions.append(f"请解释一下“{heading}”在当前知识库中的含义和作用。")
        deduped: list[str] = []
        seen: set[str] = set()
        for item in questions:
            if item not in seen:
                seen.add(item)
                deduped.append(item)
            if len(deduped) >= count:
                break
        if not deduped:
            deduped = [
                f"{kb.name} 主要覆盖了哪些主题？",
                f"使用 {kb.name} 中的信息时，应该先关注哪些内容？",
            ]
        return deduped[:count]

    def generate_sample_questions(self, kb_id: str, count: int = 10) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        files = [item for item in self.store.list_files(kb_id) if not item.is_folder]
        fallback = self._build_sample_question_fallback(kb, files, max(1, count))
        outlines = [self._load_outline(item) for item in files[:12]]
        llm_raw = self._generate_with_llm(
            system_prompt="你是知识库测试问题生成助手。只返回 JSON，格式必须是 {\"questions\": [string]}。",
            user_prompt=json.dumps(
                {
                    "knowledgeBase": kb.name,
                    "description": kb.description,
                    "count": count,
                    "files": outlines,
                },
                ensure_ascii=False,
            ),
        )
        questions = fallback
        if llm_raw:
            json_text = self._extract_json_text(llm_raw)
            if json_text:
                try:
                    candidate = json.loads(json_text)
                    values = candidate.get("questions") if isinstance(candidate, dict) else None
                    if isinstance(values, list):
                        normalized = [str(item).strip() for item in values if str(item).strip()]
                        if normalized:
                            questions = normalized[:count]
                except json.JSONDecodeError:
                    logger.warning("Sample question JSON decode failed, using fallback")

        updated = self.store.update_kb(replace(kb, sample_questions=questions, updated_at=now_iso()))
        if updated is None:
            raise KnowledgeBaseNotFoundError(kb_id)
        return {"questions": questions}

    def get_sample_questions(self, kb_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        return {"questions": list(kb.sample_questions)}

    def get_graph_labels(self, kb_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_lightrag(kb, feature="Knowledge graph")
        labels = self._run_async(self.rag_engine.get_graph_labels(kb_id))
        return {"labels": labels}

    def get_graph(self, kb_id: str, *, node_label: str = "*", max_depth: int = 2, max_nodes: int = 50) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_lightrag(kb, feature="Knowledge graph")
        graph = self._run_async(
            self.rag_engine.get_knowledge_graph(
                kb_id,
                label=node_label or "*",
                max_depth=max(1, int(max_depth)),
                max_nodes=max(10, int(max_nodes)),
            )
        )
        return graph

    def get_graph_stats(self, kb_id: str) -> dict[str, Any]:
        graph = self.get_graph(kb_id, node_label="*", max_depth=2, max_nodes=200)
        return {
            "nodeCount": len(graph.get("nodes") or []),
            "edgeCount": len(graph.get("edges") or []),
            "labels": list(graph.get("labels") or []),
            "isTruncated": bool(graph.get("isTruncated")),
        }

    def _benchmark_data_path(self, kb_id: str, benchmark_id: str) -> Path:
        return self._kb_benchmarks_dir(kb_id) / f"{benchmark_id}.jsonl"

    def _benchmark_meta_path(self, kb_id: str, benchmark_id: str) -> Path:
        return self._kb_benchmarks_dir(kb_id) / f"{benchmark_id}.meta.json"

    def _evaluation_result_path(self, kb_id: str, task_id: str) -> Path:
        return self._kb_results_dir(kb_id) / f"{task_id}.json"

    def _save_benchmark(
        self,
        kb_id: str,
        benchmark_id: str,
        *,
        name: str,
        description: str,
        questions: list[dict[str, Any]],
        created_by: str | None = None,
    ) -> dict[str, Any]:
        now = now_iso()
        data_path = self._benchmark_data_path(kb_id, benchmark_id)
        data_path.parent.mkdir(parents=True, exist_ok=True)
        lines = [json.dumps(item, ensure_ascii=False) for item in questions]
        data_path.write_text("\n".join(lines), encoding="utf-8")
        meta = {
            "id": benchmark_id,
            "benchmarkId": benchmark_id,
            "benchmark_id": benchmark_id,
            "dbId": kb_id,
            "db_id": kb_id,
            "name": name,
            "description": description,
            "questionCount": len(questions),
            "question_count": len(questions),
            "hasGoldChunks": any(bool(item.get("gold_chunk_ids")) for item in questions),
            "has_gold_chunks": any(bool(item.get("gold_chunk_ids")) for item in questions),
            "hasGoldAnswers": any(bool(item.get("gold_answer")) for item in questions),
            "has_gold_answers": any(bool(item.get("gold_answer")) for item in questions),
            "benchmarkFile": str(data_path),
            "benchmark_file": str(data_path),
            "createdBy": created_by,
            "created_by": created_by,
            "createdAt": now,
            "created_at": now,
            "updatedAt": now,
            "updated_at": now,
        }
        with self._evaluation_lock:
            self._atomic_write_json(self._benchmark_meta_path(kb_id, benchmark_id), meta)
        return meta

    def _load_benchmark_meta(self, kb_id: str, benchmark_id: str) -> dict[str, Any]:
        meta = self._read_json(self._benchmark_meta_path(kb_id, benchmark_id), {})
        if not isinstance(meta, dict) or not meta:
            raise KnowledgeBaseValidationError("Benchmark not found.")
        return meta

    def _load_benchmark_questions(self, kb_id: str, benchmark_id: str) -> list[dict[str, Any]]:
        path = self._benchmark_data_path(kb_id, benchmark_id)
        if not path.exists():
            raise KnowledgeBaseValidationError("Benchmark file is missing.")
        questions: list[dict[str, Any]] = []
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            item = json.loads(stripped)
            if isinstance(item, dict):
                questions.append(item)
        return questions

    def list_benchmarks(self, kb_id: str) -> list[dict[str, Any]]:
        kb = self.require_kb(kb_id)
        self._ensure_evaluation_supported(kb)
        result: list[dict[str, Any]] = []
        for path in sorted(self._kb_benchmarks_dir(kb_id).glob("*.meta.json"), reverse=True):
            meta = self._read_json(path, {})
            if isinstance(meta, dict) and meta:
                result.append(meta)
        result.sort(key=lambda item: str(item.get("updatedAt") or item.get("updated_at") or ""), reverse=True)
        return result

    def upload_benchmark(
        self,
        kb_id: str,
        *,
        file_content: bytes,
        filename: str,
        name: str,
        description: str,
        created_by: str | None = None,
    ) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_evaluation_supported(kb)
        if not filename.lower().endswith(".jsonl"):
            raise KnowledgeBaseValidationError("Benchmark upload only supports JSONL files.")
        try:
            content = file_content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise KnowledgeBaseValidationError("Benchmark file must be UTF-8 encoded JSONL.") from exc

        questions: list[dict[str, Any]] = []
        for line_no, line in enumerate(content.splitlines(), start=1):
            stripped = line.strip()
            if not stripped:
                continue
            try:
                payload = json.loads(stripped)
            except json.JSONDecodeError as exc:
                raise KnowledgeBaseValidationError(f"Benchmark JSONL line {line_no} is invalid: {exc}") from exc
            if not isinstance(payload, dict) or not str(payload.get("query") or "").strip():
                raise KnowledgeBaseValidationError(f"Benchmark JSONL line {line_no} is missing the 'query' field.")
            questions.append(payload)
        if not questions:
            raise KnowledgeBaseValidationError("Benchmark file does not contain any valid questions.")

        benchmark_id = _short_id("benchmark")
        return self._save_benchmark(
            kb_id,
            benchmark_id,
            name=name.strip() or Path(filename).stem or benchmark_id,
            description=description.strip(),
            questions=questions,
            created_by=created_by,
        )

    def get_benchmark_detail(
        self,
        kb_id: str,
        benchmark_id: str,
        *,
        page: int = 1,
        page_size: int = 10,
    ) -> dict[str, Any]:
        meta = self._load_benchmark_meta(kb_id, benchmark_id)
        questions = self._load_benchmark_questions(kb_id, benchmark_id)
        page = max(1, int(page))
        page_size = max(1, min(int(page_size), 100))
        start = (page - 1) * page_size
        end = start + page_size
        total = len(questions)
        total_pages = max(1, math.ceil(total / page_size))
        return {
            **meta,
            "questions": questions[start:end],
            "pagination": {
                "currentPage": page,
                "current_page": page,
                "pageSize": page_size,
                "page_size": page_size,
                "totalQuestions": total,
                "total_questions": total,
                "totalPages": total_pages,
                "total_pages": total_pages,
                "hasNext": page < total_pages,
                "hasPrev": page > 1,
            },
        }

    def delete_benchmark(self, kb_id: str, benchmark_id: str) -> bool:
        self._load_benchmark_meta(kb_id, benchmark_id)
        for path in (self._benchmark_data_path(kb_id, benchmark_id), self._benchmark_meta_path(kb_id, benchmark_id)):
            if path.exists():
                path.unlink()
        return True

    def get_benchmark_download_path(self, kb_id: str, benchmark_id: str) -> Path:
        self._load_benchmark_meta(kb_id, benchmark_id)
        path = self._benchmark_data_path(kb_id, benchmark_id)
        if not path.exists():
            raise KnowledgeBaseValidationError("Benchmark file is missing from disk.")
        return path

    def generate_benchmark(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_evaluation_supported(kb)
        count = max(1, min(int(payload.get("count") or 10), 50))
        name = self._normalize_text(payload.get("name"), field_name="name") or "自动生成评估基准"
        description = self._normalize_text(payload.get("description"), field_name="description")
        entries = self._load_vector_entries(kb_id)
        if not entries:
            raise KnowledgeBaseValidationError("Generate benchmark requires indexed Milvus chunks.")

        selected_entries = entries[: min(len(entries), count)]
        questions: list[dict[str, Any]] = []
        for item in selected_entries:
            chunk_id = str(item.get("chunkId") or "")
            content = str(item.get("content") or "").strip()
            prompt = self._generate_with_llm(
                system_prompt=(
                    "你是 RAG 评测基准生成助手。请基于给定片段生成一个可由片段直接回答的问题。"
                    "只返回 JSON，格式为 {\"query\": string, \"gold_answer\": string}。"
                ),
                user_prompt=json.dumps(
                    {
                        "knowledgeBase": kb.name,
                        "chunk": {
                            "chunkId": chunk_id,
                            "file": item.get("filename"),
                            "content": content[:2000],
                        },
                    },
                    ensure_ascii=False,
                ),
            )
            query_text = ""
            gold_answer = ""
            json_text = self._extract_json_text(prompt or "")
            if json_text:
                try:
                    candidate = json.loads(json_text)
                    if isinstance(candidate, dict):
                        query_text = str(candidate.get("query") or "").strip()
                        gold_answer = str(candidate.get("gold_answer") or "").strip()
                except json.JSONDecodeError:
                    logger.warning("Benchmark generation JSON decode failed for {}", chunk_id)
            if not query_text:
                first_line = next((line.strip() for line in content.splitlines() if line.strip()), kb.name)
                query_text = f"请根据知识库内容解释：{first_line[:40]}"
            if not gold_answer:
                gold_answer = content[:400]
            questions.append(
                {
                    "query": query_text,
                    "gold_answer": gold_answer,
                    "gold_chunk_ids": [chunk_id] if chunk_id else [],
                }
            )
            if len(questions) >= count:
                break

        benchmark_id = _short_id("benchmark")
        return self._save_benchmark(
            kb_id,
            benchmark_id,
            name=name,
            description=description,
            questions=questions,
        )

    @staticmethod
    def _retrieval_metrics(retrieved_chunks: list[dict[str, Any]], gold_chunk_ids: list[str]) -> dict[str, float]:
        if not retrieved_chunks or not gold_chunk_ids:
            return {}
        retrieved_ids = [str(item.get("chunk_id") or item.get("chunkId") or "") for item in retrieved_chunks]
        relevant = {str(item) for item in gold_chunk_ids if str(item).strip()}
        if not relevant:
            return {}
        metrics: dict[str, float] = {}
        for k in (1, 3, 5, 10):
            top_k = retrieved_ids[:k]
            if not top_k:
                metrics[f"precision@{k}"] = 0.0
                metrics[f"recall@{k}"] = 0.0
                metrics[f"f1@{k}"] = 0.0
                continue
            hit_count = len(set(top_k) & relevant)
            precision = hit_count / k
            recall = hit_count / len(relevant)
            f1 = 0.0 if precision + recall == 0 else (2 * precision * recall / (precision + recall))
            metrics[f"precision@{k}"] = precision
            metrics[f"recall@{k}"] = recall
            metrics[f"f1@{k}"] = f1
        return metrics

    @staticmethod
    def _normalize_eval_text(text: str) -> str:
        lowered = str(text or "").strip().lower()
        return re.sub(r"\s+", " ", re.sub(r"[^\w\u4e00-\u9fff]+", " ", lowered)).strip()

    def _answer_metrics(self, query: str, generated_answer: str, gold_answer: str) -> dict[str, Any]:
        if not gold_answer.strip():
            return {}
        if not generated_answer.strip():
            return {"score": 0.0, "reasoning": "未生成答案"}

        llm_raw = self._generate_with_llm(
            system_prompt=(
                "你是 RAG 评测裁判。只返回 JSON，格式为 "
                "{\"score\": 0 或 1, \"reasoning\": string}。"
            ),
            user_prompt=json.dumps(
                {
                    "query": query,
                    "goldAnswer": gold_answer,
                    "generatedAnswer": generated_answer,
                },
                ensure_ascii=False,
            ),
        )
        json_text = self._extract_json_text(llm_raw or "")
        if json_text:
            try:
                payload = json.loads(json_text)
                if isinstance(payload, dict):
                    return {
                        "score": float(payload.get("score") or 0.0),
                        "reasoning": str(payload.get("reasoning") or "").strip(),
                    }
            except json.JSONDecodeError:
                logger.warning("Answer metric JSON decode failed, falling back to heuristic scoring")

        normalized_generated = self._normalize_eval_text(generated_answer)
        normalized_gold = self._normalize_eval_text(gold_answer)
        score = 1.0 if normalized_gold and normalized_gold in normalized_generated else 0.0
        return {
            "score": score,
            "reasoning": "使用启发式匹配完成评分",
        }

    def _load_evaluation_result(self, kb_id: str, task_id: str) -> dict[str, Any]:
        payload = self._read_json(self._evaluation_result_path(kb_id, task_id), {})
        if not isinstance(payload, dict) or not payload:
            raise KnowledgeBaseValidationError("Evaluation result not found.")
        return payload

    def _save_evaluation_result(self, kb_id: str, task_id: str, payload: dict[str, Any]) -> None:
        with self._evaluation_lock:
            self._atomic_write_json(self._evaluation_result_path(kb_id, task_id), payload)

    def run_evaluation(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        self._ensure_evaluation_supported(kb)
        benchmark_id = self._normalize_text(
            payload.get("benchmarkId") or payload.get("benchmark_id"),
            required=True,
            field_name="benchmarkId",
        )
        self._load_benchmark_meta(kb_id, benchmark_id)
        task_id = _short_id("eval")
        result = {
            "taskId": task_id,
            "task_id": task_id,
            "kbId": kb_id,
            "dbId": kb_id,
            "benchmarkId": benchmark_id,
            "benchmark_id": benchmark_id,
            "status": "queued",
            "overallScore": None,
            "overall_score": None,
            "totalQuestions": 0,
            "total_questions": 0,
            "completedQuestions": 0,
            "completed_questions": 0,
            "retrievalConfig": kb.query_params.to_dict(),
            "retrieval_config": kb.query_params.to_dict(),
            "modelConfig": dict(payload.get("modelConfig") or payload.get("model_config") or {}),
            "model_config": dict(payload.get("modelConfig") or payload.get("model_config") or {}),
            "details": [],
            "metrics": {},
            "createdAt": now_iso(),
            "created_at": now_iso(),
            "updatedAt": now_iso(),
            "updated_at": now_iso(),
        }
        self._save_evaluation_result(kb_id, task_id, result)
        self._submit_background_job(self._run_evaluation_task, kb_id, task_id)
        return {"taskId": task_id, "task_id": task_id}

    def _run_evaluation_task(self, kb_id: str, task_id: str) -> None:
        result = self._load_evaluation_result(kb_id, task_id)
        benchmark_id = str(result.get("benchmarkId") or result.get("benchmark_id") or "")
        questions = self._load_benchmark_questions(kb_id, benchmark_id)
        result["status"] = "running"
        result["startedAt"] = now_iso()
        result["started_at"] = result["startedAt"]
        result["totalQuestions"] = len(questions)
        result["total_questions"] = len(questions)
        self._save_evaluation_result(kb_id, task_id, result)

        retrieval_metric_list: list[dict[str, float]] = []
        answer_metric_list: list[dict[str, Any]] = []
        detail_items: list[dict[str, Any]] = []

        try:
            retrieval_config = dict(result.get("retrievalConfig") or result.get("retrieval_config") or {})
            for index, question in enumerate(questions, start=1):
                query_text = str(question.get("query") or "").strip()
                gold_answer = str(question.get("gold_answer") or "").strip()
                gold_chunk_ids = [str(item) for item in question.get("gold_chunk_ids") or [] if str(item).strip()]

                query_result = self.query_database(
                    kb_id,
                    {
                        **retrieval_config,
                        "query": query_text,
                    },
                )
                chunks = list((query_result.get("data") or {}).get("chunks") or [])
                generated_answer = str(query_result.get("message") or "").strip()
                retrieval_metrics = self._retrieval_metrics(chunks, gold_chunk_ids)
                answer_metrics = self._answer_metrics(query_text, generated_answer, gold_answer) if gold_answer else {}

                retrieval_metric_list.append(retrieval_metrics)
                if answer_metrics:
                    answer_metric_list.append(answer_metrics)

                detail_items.append(
                    {
                        "rowId": f"{task_id}-{index}",
                        "row_id": f"{task_id}-{index}",
                        "query": query_text,
                        "goldAnswer": gold_answer,
                        "gold_answer": gold_answer,
                        "goldChunkIds": gold_chunk_ids,
                        "gold_chunk_ids": gold_chunk_ids,
                        "generatedAnswer": generated_answer,
                        "generated_answer": generated_answer,
                        "retrievedChunks": chunks,
                        "retrieved_chunks": chunks,
                        "metrics": {**retrieval_metrics, **answer_metrics},
                        "errorMessage": None,
                        "error_message": None,
                    }
                )
                result["details"] = detail_items
                result["completedQuestions"] = index
                result["completed_questions"] = index
                result["updatedAt"] = now_iso()
                result["updated_at"] = result["updatedAt"]
                self._save_evaluation_result(kb_id, task_id, result)

            aggregate_metrics: dict[str, float] = {}
            metric_keys = sorted({key for item in retrieval_metric_list for key in item.keys()})
            for key in metric_keys:
                values = [item[key] for item in retrieval_metric_list if key in item]
                if values:
                    aggregate_metrics[key] = sum(values) / len(values)
            if answer_metric_list:
                aggregate_metrics["answer_accuracy"] = sum(
                    float(item.get("score") or 0.0) for item in answer_metric_list
                ) / len(answer_metric_list)

            overall_components: list[float] = []
            for item in retrieval_metric_list:
                if item:
                    overall_components.append(sum(item.values()) / len(item))
            overall_components.extend(float(item.get("score") or 0.0) for item in answer_metric_list)
            overall_score = sum(overall_components) / len(overall_components) if overall_components else 0.0

            result["status"] = "completed"
            result["overallScore"] = overall_score
            result["overall_score"] = overall_score
            result["metrics"] = aggregate_metrics
            result["finishedAt"] = now_iso()
            result["finished_at"] = result["finishedAt"]
            result["updatedAt"] = result["finishedAt"]
            result["updated_at"] = result["finished_at"]
            self._save_evaluation_result(kb_id, task_id, result)
        except Exception as exc:
            logger.exception("Evaluation task {} failed", task_id)
            result["status"] = "failed"
            result["errorSummary"] = str(exc)
            result["error_summary"] = str(exc)
            result["finishedAt"] = now_iso()
            result["finished_at"] = result["finishedAt"]
            result["updatedAt"] = result["finishedAt"]
            result["updated_at"] = result["finished_at"]
            self._save_evaluation_result(kb_id, task_id, result)

    def get_evaluation_history(self, kb_id: str) -> list[dict[str, Any]]:
        kb = self.require_kb(kb_id)
        self._ensure_evaluation_supported(kb)
        result: list[dict[str, Any]] = []
        for path in sorted(self._kb_results_dir(kb_id).glob("*.json"), reverse=True):
            payload = self._read_json(path, {})
            if isinstance(payload, dict) and payload:
                summary = {key: value for key, value in payload.items() if key != "details"}
                result.append(summary)
        result.sort(key=lambda item: str(item.get("updatedAt") or item.get("updated_at") or ""), reverse=True)
        return result

    def get_evaluation_result(
        self,
        kb_id: str,
        task_id: str,
        *,
        page: int = 1,
        page_size: int = 20,
        error_only: bool = False,
    ) -> dict[str, Any]:
        result = self._load_evaluation_result(kb_id, task_id)
        details = list(result.get("details") or [])
        if error_only:
            details = [
                item
                for item in details
                if str(item.get("errorMessage") or item.get("error_message") or "").strip()
                or float((item.get("metrics") or {}).get("score") or 1.0) < 1.0
                or float((item.get("metrics") or {}).get("recall@1") or 1.0) < 1.0
            ]
        page = max(1, int(page))
        page_size = max(1, min(int(page_size), 100))
        start = (page - 1) * page_size
        end = start + page_size
        total = len(details)
        total_pages = max(1, math.ceil(total / page_size))
        return {
            **{key: value for key, value in result.items() if key != "details"},
            "details": details[start:end],
            "pagination": {
                "currentPage": page,
                "current_page": page,
                "pageSize": page_size,
                "page_size": page_size,
                "total": total,
                "total_pages": total_pages,
                "hasNext": page < total_pages,
                "hasPrev": page > 1,
            },
        }

    def delete_evaluation_result(self, kb_id: str, task_id: str) -> bool:
        path = self._evaluation_result_path(kb_id, task_id)
        if not path.exists():
            raise KnowledgeBaseValidationError("Evaluation result not found.")
        path.unlink()
        return True

    def get_download_path(self, kb_id: str, file_id: str, *, variant: str = "raw") -> Path:
        self.require_kb(kb_id)
        file = self._require_file(kb_id, file_id)
        candidate = file.raw_path if variant != "parsed" else (file.markdown_file or file.raw_path)
        if not candidate:
            raise KnowledgeBaseValidationError("Requested download artifact is not available.")
        path = Path(candidate)
        if not path.exists():
            raise KnowledgeBaseValidationError("Requested download artifact is missing from disk.")
        return path

    def get_file_detail(self, kb_id: str, file_id: str) -> dict[str, Any]:
        kb = self.require_kb(kb_id)
        file = self._require_file(kb_id, file_id)
        if file.is_folder:
            raise KnowledgeBaseValidationError("Folder detail preview is not supported.")

        content = ""
        if file.markdown_file and Path(file.markdown_file).exists():
            content = Path(file.markdown_file).read_text(encoding="utf-8")
        elif file.raw_path and Path(file.raw_path).exists():
            raw_bytes = Path(file.raw_path).read_bytes()
            suffix = Path(file.raw_path).suffix.lower()
            is_text_like = (
                str(file.content_type or "").startswith("text/")
                or suffix in {".md", ".txt", ".json", ".csv", ".html", ".htm", ".xml", ".yaml", ".yml"}
            )
            if is_text_like:
                content = self._decode_text(raw_bytes)

        chunks = [
            {
                "chunkId": item.get("chunkId"),
                "chunk_id": item.get("chunkId"),
                "chunkIndex": item.get("chunkIndex"),
                "chunk_index": item.get("chunkIndex"),
                "content": item.get("content"),
                "fileId": item.get("fileId"),
                "file_id": item.get("fileId"),
                "filename": item.get("filename"),
                "file_path": item.get("path") or item.get("filePath"),
                "metadata": {
                    "path": item.get("path"),
                    "filePath": item.get("filePath"),
                },
            }
            for item in self._load_vector_entries(kb.kb_id)
            if str(item.get("fileId") or "") == file.file_id
        ]
        chunks.sort(key=lambda item: int(item.get("chunkIndex") or 0))

        return {
            "file": self._serialize_file(file),
            "content": content,
            "chunks": chunks,
            "chunkCount": len(chunks) or int(file.processing_params.get("chunksCount") or 0),
        }

    def query_kb_for_agent(
        self,
        kb_id: str,
        query_text: str,
        *,
        file_name: str | None = None,
        limit: int = 6,
    ) -> dict[str, Any]:
        payload: dict[str, Any] = {"query": query_text, "topK": limit}
        if file_name:
            payload["fileName"] = file_name
        return self.query_database(kb_id, payload)

    def get_mindmap_text(self, kb_id: str) -> str:
        kb = self.require_kb(kb_id)
        if kb.mindmap is None:
            kb_map = self.generate_mindmap(kb_id)["mindmap"]
        else:
            kb_map = kb.mindmap

        lines: list[str] = []

        def _walk(node: dict[str, Any], level: int) -> None:
            content = str(node.get("content") or "").strip()
            if content:
                lines.append(f"{'  ' * level}- {content}")
            for child in node.get("children") or []:
                if isinstance(child, dict):
                    _walk(child, level + 1)

        if isinstance(kb_map, dict):
            _walk(kb_map, 0)
        return "\n".join(lines).strip()
