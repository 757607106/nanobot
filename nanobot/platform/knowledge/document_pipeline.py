"""Document parsing, indexing, and ingestion pipeline for knowledge bases.

Extracted from KnowledgeBaseService (Phase 3) to encapsulate the complete
document lifecycle: raw-file parsing → markdown conversion → LightRAG indexing
→ chunk manifest tracking.
"""

from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import re
from dataclasses import replace
from pathlib import Path
from typing import TYPE_CHECKING, Any

import chardet
from loguru import logger
from lxml import html as lxml_html
from openpyxl import load_workbook
from readability import Document as ReadabilityDocument

from nanobot.platform.knowledge.models import (
    KnowledgeBaseDefinition,
    KnowledgeDocumentStatus,
    KnowledgeFile,
    KnowledgeIngestJob,
    KnowledgeJobStatus,
    now_iso,
)
from nanobot.platform.knowledge.store import KnowledgeBaseStore
from nanobot.platform.knowledge.service import (
    KnowledgeBaseValidationError,
    KnowledgeSourceNotFoundError,
)
from nanobot.platform.knowledge.utils import (
    DEFAULT_KNOWLEDGE_CHUNK_SIZE,
    normalize_text,
    normalize_string_list,
    normalize_object,
)

if TYPE_CHECKING:
    from nanobot.platform.knowledge.artifacts import KnowledgeArtifactStore
    from nanobot.platform.knowledge.rag_engine import RAGEngine


class DocumentPipeline:
    """Handles the full document lifecycle: parse → index → chunk tracking."""

    def __init__(
        self,
        *,
        store: KnowledgeBaseStore,
        artifacts: KnowledgeArtifactStore,
        rag_engine: RAGEngine | None,
        run_async_fn: Any,
        file_storage_paths_fn: Any,
        require_kb_fn: Any,
        require_file_fn: Any,
        update_file_fn: Any,
        create_job_fn: Any,
        start_job_fn: Any,
        finish_job_fn: Any,
        submit_job_fn: Any,
        serialize_file_fn: Any,
        serialize_job_fn: Any,
        store_job_options_fn: Any,
        consume_job_options_fn: Any,
        move_file_fn: Any,
        generate_questions_fn: Any,
        resolve_vision_runtime_fn: Any = None,
    ) -> None:
        self.store = store
        self.artifacts = artifacts
        self.rag_engine = rag_engine
        self._run_async = run_async_fn
        self._file_storage_paths = file_storage_paths_fn
        self.require_kb = require_kb_fn
        self._require_file = require_file_fn
        self._update_file = update_file_fn
        self._create_job = create_job_fn
        self._start_job = start_job_fn
        self._finish_job = finish_job_fn
        self._submit_background_job = submit_job_fn
        self._serialize_file = serialize_file_fn
        self._serialize_job = serialize_job_fn
        self._store_job_options = store_job_options_fn
        self._consume_job_options = consume_job_options_fn
        self._move_file = move_file_fn
        self._generate_questions = generate_questions_fn
        self._resolve_vision_runtime = resolve_vision_runtime_fn

    # ── Text / format helpers ──────────────────────────────────────────────

    @staticmethod
    def _detect_encoding(content: bytes) -> str:
        detection = chardet.detect(content)
        return str(detection.get("encoding") or "utf-8")

    def _decode_text(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode(self._detect_encoding(content), errors="ignore")

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"[ \t]+\n", "\n", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _html_to_text(self, raw_html: str) -> tuple[str, str | None]:
        doc = ReadabilityDocument(raw_html)
        title = doc.short_title() or None
        summary_html = doc.summary(html_partial=True)
        text = lxml_html.fromstring(summary_html).text_content()
        return self._normalize_whitespace(text), title

    def _json_to_text(self, raw_json: str) -> tuple[str, list[dict[str, Any]] | None]:
        payload = json.loads(raw_json)
        if isinstance(payload, list) and payload and all(isinstance(item, dict) for item in payload):
            faq_items = []
            for item in payload:
                question = str(item.get("question") or item.get("q") or "").strip()
                answer = str(item.get("answer") or item.get("a") or "").strip()
                if question and answer:
                    faq_items.append({"question": question, "answer": answer})
            if faq_items:
                return "\n\n".join(f"Q: {item['question']}\nA: {item['answer']}" for item in faq_items), faq_items
        return json.dumps(payload, ensure_ascii=False, indent=2), None

    def _csv_to_text(self, raw_csv: str) -> tuple[str, list[dict[str, Any]] | None]:
        reader = csv.DictReader(io.StringIO(raw_csv))
        rows = list(reader)
        faq_items = []
        lines: list[str] = []
        for row in rows:
            question = str(row.get("question") or row.get("q") or "").strip()
            answer = str(row.get("answer") or row.get("a") or "").strip()
            if question and answer:
                faq_items.append({"question": question, "answer": answer})
            if row:
                lines.append(" | ".join(f"{key}: {value}" for key, value in row.items()))
        if faq_items:
            return "\n\n".join(f"Q: {item['question']}\nA: {item['answer']}" for item in faq_items), faq_items
        return "\n".join(lines), None

    @staticmethod
    def _xlsx_to_text(content: bytes) -> str:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(cell or "").strip() for cell in rows[0]]
            sections.append(f"# Sheet: {sheet.title}")
            for row in rows[1:]:
                line = " | ".join(
                    f"{header[index] or f'column_{index + 1}'}: {str(value or '').strip()}"
                    for index, value in enumerate(row)
                    if str(value or "").strip()
                )
                if line:
                    sections.append(line)
        return "\n".join(sections).strip()

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeBaseValidationError("PDF parsing requires optional dependency 'pypdf'.") from exc
        reader = PdfReader(io.BytesIO(content))
        return self._normalize_whitespace("\n\n".join(page.extract_text() or "" for page in reader.pages))

    def _extract_pdf_with_vision(
        self,
        content: bytes,
        *,
        vision_runtime: dict[str, Any],
        max_images_per_page: int = 5,
        max_image_bytes: int = 4 * 1024 * 1024,
    ) -> str:
        """Extract PDF text with Vision-based image descriptions.

        For each page, first extracts text, then extracts embedded images
        and calls the Vision model to describe them.  Descriptions are
        interleaved with page text in the output.
        """
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeBaseValidationError(
                "PDF parsing requires optional dependency 'pypdf'."
            ) from exc

        from nanobot.platform.knowledge.llm_helpers import KnowledgeLLMHelper

        reader = PdfReader(io.BytesIO(content))
        sections: list[str] = []
        total_images_described = 0

        for page_idx, page in enumerate(reader.pages):
            # 1. Extract text for this page
            page_text = (page.extract_text() or "").strip()
            if page_text:
                sections.append(page_text)

            # 2. Extract images from this page
            try:
                page_images = list(page.images)
            except Exception:
                logger.debug("Failed to extract images from page {}", page_idx + 1)
                continue

            images_on_page = 0
            for img in page_images:
                if images_on_page >= max_images_per_page:
                    break
                try:
                    img_data: bytes = img.data
                    if len(img_data) < 500:  # skip tiny artifacts
                        continue
                    if len(img_data) > max_image_bytes:
                        logger.debug(
                            "Skipping oversized image ({} bytes) on page {}",
                            len(img_data), page_idx + 1,
                        )
                        continue

                    # Determine MIME type from image name
                    img_name = str(getattr(img, "name", "") or "").lower()
                    if img_name.endswith(".jpg") or img_name.endswith(".jpeg"):
                        mime = "image/jpeg"
                    elif img_name.endswith(".png"):
                        mime = "image/png"
                    elif img_name.endswith(".gif"):
                        mime = "image/gif"
                    elif img_name.endswith(".webp"):
                        mime = "image/webp"
                    else:
                        mime = "image/png"  # default fallback

                    img_b64 = base64.b64encode(img_data).decode("ascii")
                    description = self._run_async(
                        KnowledgeLLMHelper(
                            config=None, run_async_fn=self._run_async,
                        ).describe_image_async(
                            image_base64=img_b64,
                            mime_type=mime,
                            vision_runtime=vision_runtime,
                        )
                    )

                    if description:
                        sections.append(
                            f"\n[图片描述 (第{page_idx + 1}页)]\n{description}\n"
                        )
                        images_on_page += 1
                        total_images_described += 1
                    else:
                        logger.debug(
                            "Vision model returned empty for image on page {}",
                            page_idx + 1,
                        )
                except Exception:
                    logger.exception(
                        "Failed to describe image on page {}", page_idx + 1
                    )

        logger.info(
            "PDF vision parsing complete: {} pages, {} images described",
            len(reader.pages), total_images_described,
        )
        return self._normalize_whitespace("\n\n".join(sections))

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise KnowledgeBaseValidationError("DOCX parsing requires optional dependency 'python-docx'.") from exc
        document = Document(io.BytesIO(content))
        return self._normalize_whitespace("\n".join(paragraph.text for paragraph in document.paragraphs))

    def _parse_file_content(
        self,
        *,
        title: str,
        file_name: str,
        mime_type: str | None,
        content: bytes,
        vision_runtime: dict[str, Any] | None = None,
    ) -> tuple[str, str, dict[str, Any], list[dict[str, Any]] | None]:
        suffix = Path(file_name).suffix.lower()
        parser_name = suffix.lstrip(".") or "text"
        metadata: dict[str, Any] = {}
        faq_items: list[dict[str, Any]] | None = None

        if suffix in {".txt", ".md"}:
            text = self._decode_text(content)
        elif suffix in {".html", ".htm"}:
            text, detected_title = self._html_to_text(self._decode_text(content))
            if detected_title and detected_title != title:
                metadata["detectedTitle"] = detected_title
        elif suffix == ".json":
            text, faq_items = self._json_to_text(self._decode_text(content))
        elif suffix == ".csv":
            text, faq_items = self._csv_to_text(self._decode_text(content))
        elif suffix == ".xlsx":
            text = self._xlsx_to_text(content)
        elif suffix == ".pdf":
            if vision_runtime and vision_runtime.get("model"):
                text = self._extract_pdf_with_vision(content, vision_runtime=vision_runtime)
                parser_name = "pdf+vision"
                metadata["visionModel"] = vision_runtime.get("model", "")
            else:
                text = self._extract_pdf_text(content)
        elif suffix == ".docx":
            text = self._extract_docx_text(content)
        else:
            if mime_type and mime_type.startswith("text/"):
                text = self._decode_text(content)
            else:
                raise KnowledgeBaseValidationError(
                    f"Unsupported knowledge file type: {suffix or mime_type or file_name}"
                )

        normalized = self._normalize_whitespace(text)
        if not normalized:
            raise KnowledgeBaseValidationError("Parsed knowledge document is empty.")
        return normalized, parser_name, metadata, faq_items

    # ── Single-file processors ─────────────────────────────────────────────

    def parse_single_file(self, file: KnowledgeFile) -> KnowledgeFile:
        current = self._require_file(file.kb_id, file.file_id)
        if current.is_folder:
            return current
        if not current.raw_path:
            raise KnowledgeBaseValidationError(f"Knowledge file {current.file_id} has no source file.")
        source_path = Path(current.raw_path)
        if not source_path.exists():
            raise KnowledgeBaseValidationError(f"Knowledge source file is missing: {source_path}")

        current = self._update_file(
            replace(current, status=KnowledgeDocumentStatus.PARSING, error_message=None, updated_at=now_iso())
        )
        content = source_path.read_bytes()

        # Resolve vision runtime if KB has multimodal enabled
        vision_runtime: dict[str, Any] | None = None
        try:
            kb = self.require_kb(current.kb_id)
            additional = kb.additional_params or {}
            enable_multimodal = bool(additional.get("enable_multimodal"))
            if enable_multimodal and self._resolve_vision_runtime is not None:
                vision_info = additional.get("visionInfo") or {}
                vision_runtime = self._resolve_vision_runtime(vision_info)
                if vision_runtime and not vision_runtime.get("model"):
                    vision_runtime = None  # no valid model resolved
        except Exception:
            logger.debug("Failed to resolve vision runtime for KB {}", current.kb_id)

        text, parser_name, metadata, faq_items = self._parse_file_content(
            title=current.filename,
            file_name=current.original_filename or current.filename,
            mime_type=current.content_type,
            content=content,
            vision_runtime=vision_runtime,
        )
        parsed_path = Path(current.markdown_file or self._file_storage_paths(current.kb_id, current.file_id, current.filename)[1])
        parsed_path.parent.mkdir(parents=True, exist_ok=True)
        parsed_path.write_text(text, encoding="utf-8")

        processing_params = dict(current.processing_params)
        processing_params.update(metadata)
        processing_params["parserName"] = parser_name
        processing_params["parsedAt"] = now_iso()
        if faq_items:
            processing_params["faqItems"] = faq_items
        return self._update_file(
            replace(
                current,
                markdown_file=str(parsed_path),
                status=KnowledgeDocumentStatus.PARSED,
                processing_params=processing_params,
                error_message=None,
                updated_at=now_iso(),
            )
        )

    def index_single_file(self, file: KnowledgeFile) -> KnowledgeFile:
        current = self._require_file(file.kb_id, file.file_id)
        kb = self.require_kb(current.kb_id)
        if current.is_folder:
            return current
        if current.status not in {KnowledgeDocumentStatus.PARSED, KnowledgeDocumentStatus.INDEXED}:
            raise KnowledgeBaseValidationError(
                f"Knowledge file {current.file_id} must be parsed before indexing."
            )
        if not current.markdown_file:
            raise KnowledgeBaseValidationError(f"Knowledge file {current.file_id} has not been parsed.")

        parsed_path = Path(current.markdown_file)
        if not parsed_path.exists():
            raise KnowledgeBaseValidationError(f"Parsed knowledge file is missing: {parsed_path}")

        current = self._update_file(
            replace(current, status=KnowledgeDocumentStatus.INDEXING, error_message=None, updated_at=now_iso())
        )
        text = parsed_path.read_text(encoding="utf-8")

        # Ensure RAG engine is available
        self._ensure_lightrag(kb, feature="Document indexing")

        # Index via LightRAG Core
        index_result = self._run_async(
            self.rag_engine.insert_text(
                current.kb_id,
                text,
                doc_id=current.file_id,
                file_path=current.raw_path or current.filename,
            )
        )
        if not index_result.success:
            raise KnowledgeBaseValidationError(
                index_result.error or f"Failed to index knowledge file {current.file_id}"
            )

        # Build chunk manifest for local tracking
        from nanobot.platform.knowledge.chunking.dispatcher import build_chunks
        chunk_texts = build_chunks(
            text,
            kb_params=kb.additional_params,
            file_params=current.processing_params,
            faq_items=(current.processing_params or {}).get("faqItems"),
        )
        self.artifacts.replace_chunk_entries_for_file(
            current.kb_id,
            current.file_id,
            self.artifacts.build_chunk_manifest_entries(current, chunk_texts),
        )
        processing_params = dict(current.processing_params)
        processing_params["chunksCount"] = len(chunk_texts)
        processing_params["indexedAt"] = now_iso()
        processing_params["indexBackend"] = "lightrag"
        processing_params["graphExtraction"] = True
        if index_result.track_id:
            processing_params["trackId"] = index_result.track_id
        return self._update_file(
            replace(
                current,
                status=KnowledgeDocumentStatus.INDEXED,
                processing_params=processing_params,
                error_message=None,
                updated_at=now_iso(),
            )
        )

    def _ensure_lightrag(self, kb: KnowledgeBaseDefinition, *, feature: str) -> None:
        if self.rag_engine is None:
            raise KnowledgeBaseValidationError(f"{feature} is unavailable because the RAG engine is not configured.")

    # ── Job runners ────────────────────────────────────────────────────────

    def run_parse_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                self.parse_single_file(self._require_file(job.kb_id, file_id))
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_PARSING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update parse error state for {}", file_id)
        self._finish_job(job, error_summary="; ".join(errors) if errors else None)

    def run_index_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                self.index_single_file(self._require_file(job.kb_id, file_id))
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_INDEXING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update index error state for {}", file_id)
        finished_job = self._finish_job(job, error_summary="; ".join(errors) if errors else None)
        if finished_job.status == KnowledgeJobStatus.SUCCEEDED:
            try:
                kb = self.require_kb(job.kb_id)
                if bool((kb.additional_params or {}).get("auto_generate_questions")):
                    self._generate_questions(job.kb_id, count=10)
            except Exception:
                logger.exception("Failed to auto-generate sample questions for {}", job.kb_id)

    def run_ingest_job(self, job_id: str) -> None:
        job = self.store.get_job(job_id)
        if job is None:
            return
        job = self._start_job(job)
        options = self._consume_job_options(job_id)
        auto_index = bool(options.get("auto_index"))
        index_params = self._normalize_index_params(options.get("index_params"))
        errors: list[str] = []
        for file_id in job.target_file_ids:
            try:
                parsed = self.parse_single_file(self._require_file(job.kb_id, file_id))
                if not auto_index:
                    continue
                indexed_input = parsed
                if index_params:
                    indexed_input = self._update_file(
                        replace(
                            parsed,
                            processing_params={
                                **(parsed.processing_params or {}),
                                **index_params,
                            },
                            updated_at=now_iso(),
                        )
                    )
                self.index_single_file(indexed_input)
            except Exception as exc:
                errors.append(f"{file_id}: {exc}")
                try:
                    file = self._require_file(job.kb_id, file_id)
                    self._update_file(
                        replace(
                            file,
                            status=KnowledgeDocumentStatus.ERROR_INDEXING if auto_index else KnowledgeDocumentStatus.ERROR_PARSING,
                            error_message=str(exc),
                            updated_at=now_iso(),
                        )
                    )
                except Exception:
                    logger.exception("Failed to update ingest error state for {}", file_id)
        finished_job = self._finish_job(job, error_summary="; ".join(errors) if errors else None)
        if auto_index and finished_job.status == KnowledgeJobStatus.SUCCEEDED:
            try:
                kb = self.require_kb(job.kb_id)
                if bool((kb.additional_params or {}).get("auto_generate_questions")):
                    self._generate_questions(job.kb_id, count=10)
            except Exception:
                logger.exception("Failed to auto-generate sample questions for {}", job.kb_id)

    # ── Selection / resolution helpers ─────────────────────────────────────

    def _select_target_files(self, kb_id: str, file_ids: list[str] | None = None) -> list[KnowledgeFile]:
        self.require_kb(kb_id)
        files = self.store.list_files(kb_id)
        if file_ids:
            wanted = set(normalize_string_list(file_ids, field_name="fileIds"))
            selected = [item for item in files if item.file_id in wanted and not item.is_folder]
        else:
            selected = [item for item in files if not item.is_folder]
        if not selected:
            raise KnowledgeBaseValidationError("No knowledge files matched the requested operation.")
        return selected

    def _resolve_existing_files(
        self,
        kb_id: str,
        candidates: list[str],
    ) -> tuple[list[KnowledgeFile], list[str]]:
        self.require_kb(kb_id)
        files = self.store.list_files(kb_id)
        resolved: list[KnowledgeFile] = []
        seen: set[str] = set()
        missing: list[str] = []
        for raw_candidate in candidates:
            candidate = str(raw_candidate or "").strip()
            if not candidate:
                continue
            matched: KnowledgeFile | None = None
            for item in files:
                tokens = {
                    item.file_id,
                    item.filename,
                    item.path,
                    item.raw_path or "",
                    Path(item.raw_path).name if item.raw_path else "",
                }
                if candidate in tokens:
                    matched = item
                    break
            if matched is None:
                missing.append(candidate)
                continue
            if matched.file_id in seen or matched.is_folder:
                continue
            seen.add(matched.file_id)
            resolved.append(matched)
        return resolved, missing

    @staticmethod
    def _extract_requested_file_ids(payload: dict[str, Any] | None) -> list[str] | None:
        if not isinstance(payload, dict):
            return None
        for key in ("fileIds", "file_ids", "docIds", "doc_ids"):
            value = payload.get(key)
            if isinstance(value, list):
                return normalize_string_list(value, field_name=key)
        return None

    @staticmethod
    def _normalize_index_params(payload: Any) -> dict[str, Any]:
        if not isinstance(payload, dict):
            return {}
        normalized: dict[str, Any] = {}
        raw_chunk_size = payload.get("chunk_size") if "chunk_size" in payload else payload.get("chunkSize")
        if raw_chunk_size is not None:
            normalized["chunk_size"] = max(200, int(raw_chunk_size))
        raw_chunk_overlap = payload.get("chunk_overlap") if "chunk_overlap" in payload else payload.get("chunkOverlap")
        if raw_chunk_overlap is not None:
            overlap = max(0, int(raw_chunk_overlap))
            chunk_size = int(
                normalized.get("chunk_size")
                or payload.get("chunk_size")
                or payload.get("chunkSize")
                or DEFAULT_KNOWLEDGE_CHUNK_SIZE
            )
            normalized["chunk_overlap"] = min(overlap, max(0, chunk_size - 1))
        chunk_preset = normalize_text(
            payload.get("chunk_preset_id") if "chunk_preset_id" in payload else payload.get("chunkPresetId"),
            field_name="chunkPresetId",
        )
        if chunk_preset:
            normalized["chunk_preset_id"] = chunk_preset
        qa_separator = normalize_text(
            payload.get("qa_separator") if "qa_separator" in payload else payload.get("qaSeparator"),
            field_name="qaSeparator",
        )
        if qa_separator:
            normalized["qa_separator"] = qa_separator
        return normalized

    # ── Public API ─────────────────────────────────────────────────────────

    def parse_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self.require_kb(kb_id)
        requested_file_ids = self._extract_requested_file_ids(payload)
        if not requested_file_ids:
            raise KnowledgeBaseValidationError("fileIds is required for parse operations.")
        selected = self._select_target_files(kb_id, requested_file_ids)
        job = self._create_job(kb_id, "parse", [item.file_id for item in selected])
        self._submit_background_job(self.run_parse_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }

    def index_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self.require_kb(kb_id)
        requested_file_ids = self._extract_requested_file_ids(payload)
        if not requested_file_ids:
            raise KnowledgeBaseValidationError("fileIds is required for index operations.")
        selected = self._select_target_files(kb_id, requested_file_ids)
        index_params = self._normalize_index_params(payload.get("params"))
        if index_params:
            refreshed_selected: list[KnowledgeFile] = []
            for item in selected:
                if item.is_folder:
                    refreshed_selected.append(item)
                    continue
                persisted = self._update_file(
                    replace(
                        item,
                        processing_params={
                            **(item.processing_params or {}),
                            **index_params,
                        },
                        updated_at=now_iso(),
                    )
                )
                refreshed_selected.append(persisted)
            selected = refreshed_selected
        job = self._create_job(kb_id, "index", [item.file_id for item in selected])
        self._submit_background_job(self.run_index_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }

    def ingest_files(self, kb_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self.require_kb(kb_id)
        params = normalize_object(payload.get("params"), field_name="params")
        requested_file_ids = self._extract_requested_file_ids(payload)
        if requested_file_ids:
            selected = self._select_target_files(kb_id, requested_file_ids)
        else:
            items = normalize_string_list(payload.get("items"), field_name="items")
            if not items:
                raise KnowledgeBaseValidationError("items is required for ingest operations.")
            selected, missing = self._resolve_existing_files(kb_id, items)
            if missing:
                raise KnowledgeBaseValidationError(f"Unknown uploaded items: {', '.join(missing)}.")
            if not selected:
                raise KnowledgeBaseValidationError("No uploaded items matched the requested ingest payload.")

        parent_id = (
            normalize_text(params.get("parentId"), field_name="parentId")
            or normalize_text(params.get("parent_id"), field_name="parent_id")
            or None
        )
        if parent_id:
            refreshed: list[KnowledgeFile] = []
            for item in selected:
                moved = self._move_file(
                    kb_id,
                    {
                        "fileId": item.file_id,
                        "parentId": parent_id,
                    },
                )
                refreshed.append(self._require_file(kb_id, str(moved.get("fileId") or item.file_id)))
            selected = refreshed

        job = self._create_job(kb_id, "ingest", [item.file_id for item in selected])
        self._store_job_options(
            job.job_id,
            {
                "auto_index": bool(params.get("auto_index") or params.get("autoIndex")),
                "index_params": params,
            },
        )
        self._submit_background_job(self.run_ingest_job, job.job_id)
        return {
            "job": self._serialize_job(job),
            "items": [self._serialize_file(item) for item in selected],
        }
